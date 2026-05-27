"""Site Visit Requests router — multi-stage e-approval workflow for the
"بدل موقع" (site-visit allowance) form.

Flow:
  user submits → pending_head
  → head signs   → pending_supervisor
  → supervisor signs → pending_director
  → director signs   → approved

Permissions:
  - submit_site_visit       : create a new request
  - sign_as_head            : sign at the head-of-department stage
  - sign_as_supervisor      : sign at the maintenance-supervisor stage
  - sign_as_director        : sign at the department-director stage
  - view_all_site_visits    : list ALL requests (otherwise sees only own +
                              pending-for-me)
"""
import io
import json
import logging
import os
import secrets
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import or_, select
from sqlalchemy.ext.asyncio import AsyncSession

from core.auth import AccessTokenError, decode_access_token
from core.database import get_db
from models.auth import User
from models.site_visit_request import SiteVisitRequest
from models.user_roles import User_roles

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/site-visits", tags=["site-visits"])


# ---------- Schemas ----------
class SiteVisitRowItem(BaseModel):
    date: Optional[str] = None
    mosque: Optional[str] = None
    description: Optional[str] = None
    distance: Optional[str] = None
    duration: Optional[str] = None
    signature: Optional[str] = None
    extra: Optional[dict] = None


class CreateSiteVisitRequest(BaseModel):
    # `owner_name` is the recipient/submitter name TYPED INTO THE FORM by the
    # user (the "اسم المستلم" input). It is intentionally separate from the
    # account's display name (`users.name`) — the form's value is what gets
    # printed in the monthly Word export and the PDF.
    owner_name: Optional[str] = None
    civil_id: Optional[str] = None
    job_title: Optional[str] = None
    month: Optional[int] = None
    year: Optional[int] = None
    area: Optional[str] = None
    reason: Optional[str] = None
    rows: List[dict] = Field(default_factory=list)
    # Optional: web path (e.g. "/uploads/site-visit-attendance/<name>.jpg")
    # returned by POST /upload-attendance and passed back here at create time.
    attendance_attachment: Optional[str] = None


class SignSiteVisitRequest(BaseModel):
    request_id: int
    # signature_data_url is kept optional for backward compatibility with older
    # clients but is IGNORED. The backend now stamps the signer's full name
    # (or email fallback) as the signature text — the user must print the
    # form and stamp it manually after approval.
    signature_data_url: Optional[str] = None


class AuditSiteVisitRequest(BaseModel):
    """Body for POST /audit — auditor approves or rejects the request.

    `decision`:
      - "approve" → moves status from pending_audit (or rejected_audit) to
        pending_head, kicking off the 3-stage signing chain.
      - "reject"  → moves status to rejected_audit. `note` is REQUIRED in
        this case so the submitter knows what to fix.
    """

    request_id: int
    decision: str  # "approve" or "reject"
    note: Optional[str] = None


class BulkSignRequest(BaseModel):
    request_ids: List[int] = Field(default_factory=list)


class ExportPdfsZipRequest(BaseModel):
    month: int
    year: int
    # When True, the ZIP includes ALL requests for the month (pending /
    # approved / rejected). Default False = approved only (legacy behavior).
    include_unapproved: bool = False
    # Optional named scope. Recognized values:
    #   - "audit_approved" — requests that passed the audit stage; i.e.
    #       status NOT IN ('pending_audit', 'rejected_audit'). Used by the
    #       "المعتمد من التدقيق فقط" option in the frontend export-scope
    #       picker. When set, this OVERRIDES `include_unapproved`.
    # Other values (or None) fall back to the legacy
    # `include_unapproved` flag.
    scope: Optional[str] = None
    # When non-empty, the export filters by `id IN request_ids` and the
    # `month` / `year` / `include_unapproved` / `scope` fields are IGNORED.
    # This is used by the "المحددة فقط" option in the frontend export-scope
    # picker so the admin can ZIP exactly the requests they ticked via
    # checkbox, regardless of their month/year/status.
    request_ids: Optional[List[int]] = None


class OverrideApproverNames(BaseModel):
    """Optional per-batch override of the typed approver names that appear
    above each signature line in the printed PDF. When a field is None or
    empty, the original `*_signed_by_name` from the database is used."""
    head: Optional[str] = None
    supervisor: Optional[str] = None
    director: Optional[str] = None


class BatchReadyPrintRequest(BaseModel):
    """Payload for /export-ready-print-zip — admin selects N request IDs
    and gets back a single ZIP containing one PDF per request, ready for
    physical printing + manual stamping. Replaces the previous frontend
    `window.open` loop which was blocked by browser popup-blockers when
    more than 1-2 forms were opened in rapid succession.

    `override_names` (optional): dict with `head` / `supervisor` / `director`
    text fields. When provided and non-empty, each value REPLACES the
    corresponding `*_signed_by_name` in the rendered PDF only — the database
    is NOT modified. This lets admins stamp custom Arabic titles like
    "م. أحمد الجاسم" or "د. فهد الشمري" on bulk-printed forms without
    creating per-user accounts.
    """
    request_ids: List[int] = Field(default_factory=list)
    override_names: Optional[OverrideApproverNames] = None


class ExportDocxRequest(BaseModel):
    month: int
    year: int
    # NOTE: default is now True so the Word export shows EVERY request for the
    # selected month/year (pending + approved + rejected). The user reported
    # that the previous default (`False` → approved-only) made the button
    # appear empty whenever there were no fully-approved requests yet.
    include_unapproved: bool = True
    # Optional named scope. Recognized values:
    #   - "audit_approved" — requests that passed the audit stage; i.e.
    #       status NOT IN ('pending_audit', 'rejected_audit'). Used by the
    #       "المعتمد من التدقيق فقط" option in the frontend export-scope
    #       picker. When set, this OVERRIDES `include_unapproved`.
    # Other values (or None) fall back to the legacy
    # `include_unapproved` flag.
    scope: Optional[str] = None
    # When non-empty, the export filters by `id IN request_ids` and the
    # `month` / `year` / `include_unapproved` / `scope` fields are IGNORED.
    # Used by the "المحددة فقط" option in the frontend export-scope picker
    # so the admin can build a Word file containing exactly the rows they
    # ticked, regardless of their month/year/status. The output filename in
    # this case still uses the supplied month/year for clarity.
    request_ids: Optional[List[int]] = None


class SiteVisitRead(BaseModel):
    id: int
    owner_id: Optional[str]
    owner_name: Optional[str]
    civil_id: Optional[str]
    job_title: Optional[str]
    month: Optional[int]
    year: Optional[int]
    area: Optional[str]
    reason: Optional[str]
    rows: List[dict] = Field(default_factory=list)
    head_signature: Optional[str] = None
    supervisor_signature: Optional[str] = None
    director_signature: Optional[str] = None
    head_signed_at: Optional[datetime] = None
    supervisor_signed_at: Optional[datetime] = None
    director_signed_at: Optional[datetime] = None
    head_signed_by_name: Optional[str] = None
    supervisor_signed_by_name: Optional[str] = None
    director_signed_by_name: Optional[str] = None
    # Audit stage (NEW). Filled when a user with `audit_site_visit`
    # permission approves or rejects the request at the initial
    # `pending_audit` stage. `audit_note` is mandatory for rejections.
    audited_by_name: Optional[str] = None
    audited_at: Optional[datetime] = None
    audit_note: Optional[str] = None
    # Timestamp of the most recent edit performed by the original submitter
    # AFTER the request had been rejected by the auditor. Allows the auditor
    # UI to show a "🔄 تم التعديل بعد الرفض" badge so re-submissions are
    # visually distinguished from brand-new requests.
    edited_after_audit_at: Optional[datetime] = None
    status: str
    created_at: Optional[datetime] = None
    updated_at: Optional[datetime] = None
    visit_count: int = 0
    # Display name of the user account that actually submitted the request
    # (resolved from `owner_id` -> users.name). May differ from `owner_name`
    # which is the recipient name TYPED INTO THE FORM. Frontend uses this to
    # show "تم الإرسال للاعتماد بواسطة: ..." when the two names differ.
    submitted_by_name: Optional[str] = None
    # Public web path (e.g. "/uploads/site-visit-attendance/req-12-abc.jpg")
    # served by main.py's StaticFiles mount. Frontend opens this in a new tab.
    attendance_attachment: Optional[str] = None


# ---------- Helpers ----------
async def _get_user(request: Request, db: AsyncSession) -> Optional[dict]:
    auth = request.headers.get("authorization", "")
    if not auth.startswith("Bearer "):
        return None
    try:
        payload = decode_access_token(auth[7:])
    except AccessTokenError:
        return None
    uid = payload.get("sub")
    if not uid:
        return None
    role = payload.get("role", "user")
    # Prefer username/display name. NEVER fall back to email — the form
    # must always show a human name (or the user id at worst).
    name = payload.get("username") or payload.get("name")
    try:
        res = await db.execute(select(User).where(User.id == uid))
        u = res.scalar_one_or_none()
        if u:
            role = u.role or role
            name = name or getattr(u, "name", None)
    except Exception:
        pass
    return {"id": uid, "role": role, "name": name}


async def _has_perm(db: AsyncSession, user: Optional[dict], keys: List[str]) -> bool:
    if not user:
        return False
    if user.get("role") == "owner":
        return True
    role = user.get("role", "")
    granted = False
    try:
        r = await db.execute(select(User_roles).where(User_roles.value == role))
        ro = r.scalar_one_or_none()
        if ro and ro.permissions:
            perms = json.loads(ro.permissions) if isinstance(ro.permissions, str) else ro.permissions
            if isinstance(perms, dict):
                granted = any(bool(perms.get(k, False)) for k in keys)
            elif isinstance(perms, list):
                granted = any(k in perms for k in keys)
    except Exception as e:
        logger.warning(f"role perm check failed: {e}")
    # Custom user overrides
    try:
        ur = await db.execute(select(User).where(User.id == user["id"]))
        u = ur.scalar_one_or_none()
        if u and u.custom_permissions:
            cp = json.loads(u.custom_permissions) if isinstance(u.custom_permissions, str) else u.custom_permissions
            if isinstance(cp, dict):
                if any(k in cp and bool(cp[k]) for k in keys):
                    return True
                if any(k in cp and not bool(cp[k]) for k in keys) and not granted:
                    return False
    except Exception:
        pass
    return granted


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _serialize(item: SiteVisitRequest, submitted_by_name: Optional[str] = None) -> SiteVisitRead:
    rows: List[dict] = []
    if item.rows:
        try:
            parsed = json.loads(item.rows)
            if isinstance(parsed, list):
                rows = parsed
        except Exception:
            rows = []
    visit_count = sum(
        1
        for r in rows
        if isinstance(r, dict) and (str(r.get("date") or "").strip() or str(r.get("description") or "").strip())
    )
    return SiteVisitRead(
        id=item.id,
        owner_id=item.owner_id,
        owner_name=item.owner_name,
        civil_id=item.civil_id,
        job_title=item.job_title,
        month=item.month,
        year=item.year,
        area=item.area,
        reason=item.reason,
        rows=rows,
        head_signature=item.head_signature,
        supervisor_signature=item.supervisor_signature,
        director_signature=item.director_signature,
        head_signed_at=item.head_signed_at,
        supervisor_signed_at=item.supervisor_signed_at,
        director_signed_at=item.director_signed_at,
        head_signed_by_name=item.head_signed_by_name,
        supervisor_signed_by_name=item.supervisor_signed_by_name,
        director_signed_by_name=item.director_signed_by_name,
        audited_by_name=getattr(item, "audited_by_name", None),
        audited_at=getattr(item, "audited_at", None),
        audit_note=getattr(item, "audit_note", None),
        edited_after_audit_at=getattr(item, "edited_after_audit_at", None),
        status=item.status,
        created_at=item.created_at,
        updated_at=item.updated_at,
        visit_count=visit_count,
        submitted_by_name=submitted_by_name,
        attendance_attachment=item.attendance_attachment,
    )


async def _resolve_submitter_name(db: AsyncSession, owner_id: Optional[str]) -> Optional[str]:
    """Look up the display name of the user account that submitted the request.

    Returns the user's `name` field, or None if the user no longer exists or
    `owner_id` is None. Best-effort — never raises.

    NOTE: The user model lives in `models.auth` (already imported at module
    top as `User`). A previous version of this helper imported it from
    `models.user`, which doesn't exist — the resulting `ModuleNotFoundError`
    was swallowed by the `try/except`, causing `submitted_by_name` to always
    come back as `None` and the "تم الإرسال للاعتماد بواسطة" row to never
    render. We now reuse the top-level `User` directly.
    """
    if not owner_id:
        return None
    try:
        res = await db.execute(select(User).where(User.id == owner_id))
        u = res.scalars().first()
        if u:
            name = (getattr(u, "name", None) or "").strip()
            return name or None
    except Exception as exc:  # noqa: BLE001
        logger.warning(f"_resolve_submitter_name failed for owner_id={owner_id}: {exc}")
        return None
    return None


async def _notify_users_with_perm(
    db: AsyncSession,
    perm_key: str,
    title: str,
    body: str,
    request_id: int,
    exclude_user_id: Optional[str] = None,
) -> None:
    """Find all users who have `perm_key` (via role or custom override) and send
    them in-app + WS + Web Push notifications. Best-effort; never raises."""
    try:
        # 1. Role values that grant this perm
        role_values: set[str] = set()
        rres = await db.execute(select(User_roles))
        for ro in rres.scalars().all():
            try:
                if not ro.permissions:
                    continue
                perms = json.loads(ro.permissions) if isinstance(ro.permissions, str) else ro.permissions
                if isinstance(perms, dict) and perms.get(perm_key):
                    role_values.add(ro.value)
                elif isinstance(perms, list) and perm_key in perms:
                    role_values.add(ro.value)
            except Exception:
                continue
        role_values.add("owner")  # owners always pass

        # 2. Users with one of those roles OR with custom_permissions granting this key
        target_ids: set[str] = set()
        ures = await db.execute(select(User))
        for u in ures.scalars().all():
            if u.role in role_values:
                target_ids.add(str(u.id))
                continue
            if u.custom_permissions:
                try:
                    cp = json.loads(u.custom_permissions) if isinstance(u.custom_permissions, str) else u.custom_permissions
                    if isinstance(cp, dict) and cp.get(perm_key):
                        target_ids.add(str(u.id))
                except Exception:
                    continue

        if exclude_user_id:
            target_ids.discard(str(exclude_user_id))

        if not target_ids:
            return

        # 3. Persist in-app notifications
        url = f"/site-visit-requests/{request_id}"
        try:
            from models.notifications import Notifications
            for uid in target_ids:
                db.add(
                    Notifications(
                        user_id=uid,
                        type="site_visit_request",
                        message=f"{title} — {body}",
                        report_id=0,
                        is_read=False,
                    )
                )
            await db.commit()
        except Exception as e:
            logger.warning(f"persist site-visit notifs failed: {e}")
            try:
                await db.rollback()
            except Exception:
                pass

        # 4. WS push
        try:
            from services.ws_notifications import ws_notify_users

            await ws_notify_users(
                list(target_ids),
                "site_visit_request",
                f"{title} — {body}",
                0,
                extra={"site_visit_id": request_id, "url": url},
            )
        except Exception as e:
            logger.debug(f"WS notify failed (non-critical): {e}")

        # 5. Web Push
        try:
            from services.web_push_service import send_push_to_users

            await send_push_to_users(
                db,
                list(target_ids),
                title,
                body,
                notification_type="site_visit_request",
                url=url,
            )
        except Exception as e:
            logger.debug(f"web push failed (non-critical): {e}")
    except Exception as e:
        logger.warning(f"_notify_users_with_perm failed: {e}")


# ---------- Routes ----------
@router.post("/create", response_model=SiteVisitRead)
async def create_site_visit(
    data: CreateSiteVisitRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")
    if not await _has_perm(db, user, ["submit_site_visit"]):
        raise HTTPException(status_code=403, detail="ليس لديك صلاحية إرسال طلب زيارة ميدانية")

    # Validate at least one row has data (date or description)
    has_data = False
    for r in data.rows or []:
        if not isinstance(r, dict):
            continue
        if str(r.get("date") or "").strip() or str(r.get("description") or "").strip():
            has_data = True
            break
    if not has_data:
        raise HTTPException(status_code=400, detail="يجب تعبئة صف واحد على الأقل (التاريخ أو الوصف)")

    # `owner_name` MUST come from the form's "اسم المستلم" input (data.owner_name),
    # NOT from the account's display name. Falling back to the account name only
    # if the form value is empty (legacy clients that don't send the field).
    form_owner_name = (data.owner_name or "").strip() or None
    # Carry over a previously-uploaded attendance image (if the form sent one).
    # The form is expected to call POST /upload-attendance FIRST, get back the
    # web path, and pass it back here so the upload survives until the row exists.
    initial_attendance = (getattr(data, "attendance_attachment", None) or "").strip() or None
    item = SiteVisitRequest(
        owner_id=user["id"],
        owner_name=form_owner_name or user.get("name"),
        civil_id=(data.civil_id or "").strip() or None,
        job_title=(data.job_title or "").strip() or None,
        month=data.month,
        year=data.year,
        area=(data.area or "").strip() or None,
        reason=(data.reason or "").strip() or None,
        rows=json.dumps(data.rows or [], ensure_ascii=False),
        attendance_attachment=initial_attendance,
        # Initial status is `pending_audit` — the request must be audited
        # by a user with `audit_site_visit` permission BEFORE it enters
        # the 3-stage signing chain (head → supervisor → director).
        status="pending_audit",
        created_at=_now(),
        updated_at=_now(),
    )
    db.add(item)
    await db.commit()
    await db.refresh(item)
    logger.info(f"User {user['id']} created site_visit_request {item.id}")

    title = "طلب زيارة ميدانية جديد بحاجة لتدقيق"
    body = f"بحاجة للتدقيق — مقدم الطلب: {user.get('name') or user['id']}"
    await _notify_users_with_perm(db, "audit_site_visit", title, body, item.id, exclude_user_id=user["id"])

    submitter = (user.get("name") or "").strip() or None
    return _serialize(item, submitted_by_name=submitter)


@router.get("/list", response_model=List[SiteVisitRead])
async def list_site_visits(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    can_view_all = await _has_perm(db, user, ["view_all_site_visits"])
    can_audit = await _has_perm(db, user, ["audit_site_visit"])
    can_sign_head = await _has_perm(db, user, ["sign_as_head"])
    can_sign_supervisor = await _has_perm(db, user, ["sign_as_supervisor"])
    can_sign_director = await _has_perm(db, user, ["sign_as_director"])
    can_submit = await _has_perm(db, user, ["submit_site_visit"])

    if not (
        can_view_all
        or can_audit
        or can_sign_head
        or can_sign_supervisor
        or can_sign_director
        or can_submit
    ):
        raise HTTPException(status_code=403, detail="ليس لديك صلاحية الوصول لطلبات الاعتماد")

    q = select(SiteVisitRequest).order_by(SiteVisitRequest.id.desc())
    if not can_view_all:
        clauses = [SiteVisitRequest.owner_id == user["id"]]
        if can_audit:
            # Auditors see pending, previously-rejected, AND pending_head
            # requests (the latter so they can re-audit BEFORE the head of
            # department signs). The frontend hides the "إعادة التدقيق"
            # button once head_signed_at is set, but loading the row
            # is still safe — the backend gates the actual mutation.
            clauses.append(SiteVisitRequest.status == "pending_audit")
            clauses.append(SiteVisitRequest.status == "rejected_audit")
            clauses.append(SiteVisitRequest.status == "pending_head")
        if can_sign_head:
            clauses.append(SiteVisitRequest.status == "pending_head")
        if can_sign_supervisor:
            clauses.append(SiteVisitRequest.status == "pending_supervisor")
        if can_sign_director:
            clauses.append(SiteVisitRequest.status == "pending_director")
        q = q.where(or_(*clauses))

    res = await db.execute(q)
    items = res.scalars().all()
    # Bulk-resolve submitter names for all items in one query (avoid N+1).
    # Reuse the top-level `User` (from models.auth) — a previous version
    # imported `from models.user import User`, which doesn't exist; the
    # ModuleNotFoundError was silently swallowed and `submitted_by_name` was
    # always None on the list endpoint, hiding the "تم الإرسال للاعتماد
    # بواسطة" row in the details dialog.
    owner_ids = list({it.owner_id for it in items if it.owner_id})
    name_map: dict[str, str] = {}
    if owner_ids:
        try:
            ures = await db.execute(select(User).where(User.id.in_(owner_ids)))
            for u in ures.scalars().all():
                nm = (getattr(u, "name", None) or "").strip()
                if nm:
                    name_map[u.id] = nm
        except Exception as exc:  # noqa: BLE001
            logger.warning(f"site-visits list: bulk submitter-name lookup failed: {exc}")
            name_map = {}
    return [_serialize(it, submitted_by_name=name_map.get(it.owner_id) if it.owner_id else None) for it in items]


@router.get("/signers")
async def get_signers(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Return three lists of users who can sign as head / supervisor /
    director, for use as dropdown options in the "ready-to-print" modal.

    Each list contains users whose role grants the corresponding permission
    OR whose `custom_permissions` JSON explicitly grants it. Users with role
    `owner`, `admin`, or `superadmin` are EXCLUDED from all three lists per
    user request — even if their role technically grants the permission —
    because those accounts are platform/management identities and must not
    appear as default printable signers on the official paper form.

    Response shape::

        {
          "heads":       [{"id": str, "name": str, "username": str}, ...],
          "supervisors": [...],
          "directors":   [...]
        }

    Each list is sorted alphabetically by display name (case-insensitive)
    for a deterministic UX. Users with empty/null `name` fall back to
    `username` as the display value.
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="غير مصرح")

    # Roles that are management-only and must NEVER appear as printable
    # signers, regardless of which permissions their role definition holds.
    EXCLUDED_ROLES = {"owner", "admin", "superadmin"}

    perm_keys = ["sign_as_head", "sign_as_supervisor", "sign_as_director"]

    # 1. Build set of role.values that grant each permission.
    perm_to_roles: dict[str, set[str]] = {k: set() for k in perm_keys}
    try:
        rres = await db.execute(select(User_roles))
        for ro in rres.scalars().all():
            try:
                perms = (
                    json.loads(ro.permissions)
                    if isinstance(ro.permissions, str)
                    else (ro.permissions or {})
                )
            except Exception:
                perms = {}
            if isinstance(perms, dict):
                for pkey in perm_keys:
                    if bool(perms.get(pkey, False)):
                        perm_to_roles[pkey].add(ro.value)
            elif isinstance(perms, list):
                for pkey in perm_keys:
                    if pkey in perms:
                        perm_to_roles[pkey].add(ro.value)
    except Exception as e:
        logger.warning(f"signers: role load failed: {e}")

    # 2. Load all users.
    try:
        ures = await db.execute(select(User))
        all_users = list(ures.scalars().all())
    except Exception as e:
        logger.warning(f"signers: user load failed: {e}")
        all_users = []

    def user_grants(u: User, pkey: str) -> bool:
        # Custom override has highest priority (grant or explicit revoke).
        if u.custom_permissions:
            try:
                cp = (
                    json.loads(u.custom_permissions)
                    if isinstance(u.custom_permissions, str)
                    else u.custom_permissions
                )
                if isinstance(cp, dict) and pkey in cp:
                    return bool(cp[pkey])
                if isinstance(cp, list) and pkey in cp:
                    return True
            except Exception:
                pass
        # Role-based grant.
        return (u.role or "") in perm_to_roles.get(pkey, set())

    def user_to_dict(u: User) -> dict:
        nm = (u.name or "").strip()
        un = (u.username or "").strip() if hasattr(u, "username") else ""
        display = nm or un or (u.id or "")
        return {
            "id": str(u.id) if u.id is not None else "",
            "name": display,
            "username": un,
        }

    def collect(pkey: str) -> list[dict]:
        out: list[dict] = []
        seen: set[str] = set()
        for u in all_users:
            # Hard-exclude management roles.
            if (u.role or "").lower() in EXCLUDED_ROLES:
                continue
            if not user_grants(u, pkey):
                continue
            d = user_to_dict(u)
            if not d["name"] or d["id"] in seen:
                continue
            seen.add(d["id"])
            out.append(d)
        out.sort(key=lambda x: x["name"].lower())
        return out

    return {
        "heads": collect("sign_as_head"),
        "supervisors": collect("sign_as_supervisor"),
        "directors": collect("sign_as_director"),
    }


@router.get("/default-approvers")
async def get_default_approvers(
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Return the default approver names for each signature stage of the
    site-visit form. Used by the front-end "ready-to-print" mode to pre-fill
    the three signature boxes with the names of the users who currently hold
    each signing permission, so the form can be printed and stamped manually.

    Resolution order for each stage's name:
      1. Pick a user whose role has the relevant permission, OR whose
         `custom_permissions` JSON explicitly grants it.
      2. Among matching users, prefer admin/owner; otherwise pick the
         alphabetically-earliest non-empty `name` (deterministic).
      3. Fall back to "" if nobody holds the permission.

    Response shape:
        { "head_name": str, "supervisor_name": str, "director_name": str }
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="غير مصرح")

    # 1. Build set of role.values that grant each permission.
    perm_to_roles: dict[str, set[str]] = {
        "sign_as_head": set(),
        "sign_as_supervisor": set(),
        "sign_as_director": set(),
    }
    try:
        rres = await db.execute(select(User_roles))
        for ro in rres.scalars().all():
            try:
                perms = (
                    json.loads(ro.permissions)
                    if isinstance(ro.permissions, str)
                    else (ro.permissions or {})
                )
            except Exception:
                perms = {}
            if isinstance(perms, dict):
                for pkey in perm_to_roles.keys():
                    if bool(perms.get(pkey, False)):
                        perm_to_roles[pkey].add(ro.value)
            elif isinstance(perms, list):
                for pkey in perm_to_roles.keys():
                    if pkey in perms:
                        perm_to_roles[pkey].add(ro.value)
    except Exception as e:
        logger.warning(f"default-approvers: role load failed: {e}")

    # 2. Load all users once and resolve per stage.
    try:
        ures = await db.execute(select(User))
        all_users = list(ures.scalars().all())
    except Exception as e:
        logger.warning(f"default-approvers: user load failed: {e}")
        all_users = []

    def user_has_perm(u: User, pkey: str) -> bool:
        # Owner always has everything.
        if (u.role or "").lower() == "owner":
            return True
        # Custom override (positive grant or explicit denial).
        if u.custom_permissions:
            try:
                cp = (
                    json.loads(u.custom_permissions)
                    if isinstance(u.custom_permissions, str)
                    else u.custom_permissions
                )
                if isinstance(cp, dict) and pkey in cp:
                    return bool(cp[pkey])
                if isinstance(cp, list) and pkey in cp:
                    return True
            except Exception:
                pass
        # Role-based grant.
        return (u.role or "") in perm_to_roles.get(pkey, set())

    def pick_name(pkey: str) -> str:
        candidates = [u for u in all_users if user_has_perm(u, pkey)]
        if not candidates:
            return ""
        # Prefer admin / owner; then alphabetical by name.
        def sort_key(u: User) -> tuple:
            role = (u.role or "").lower()
            role_rank = 0 if role in ("owner", "admin") else 1
            return (role_rank, (u.name or "").strip().lower(), u.id or "")
        candidates.sort(key=sort_key)
        for u in candidates:
            nm = (u.name or "").strip()
            if nm:
                return nm
        return ""

    return {
        "head_name": pick_name("sign_as_head"),
        "supervisor_name": pick_name("sign_as_supervisor"),
        "director_name": pick_name("sign_as_director"),
    }


@router.get("/{request_id}", response_model=SiteVisitRead)
async def get_site_visit(
    request_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    res = await db.execute(select(SiteVisitRequest).where(SiteVisitRequest.id == request_id))
    item = res.scalar_one_or_none()
    if not item:
        raise HTTPException(status_code=404, detail="الطلب غير موجود")

    if user.get("role") not in ("admin", "owner"):
        is_owner = item.owner_id == user["id"]
        can_view_all = await _has_perm(db, user, ["view_all_site_visits"])
        # Whoever can act at the current stage can view. `pending_audit`
        # and `rejected_audit` are visible to anyone with `audit_site_visit`.
        stage_perm_map = {
            "pending_audit": "audit_site_visit",
            "rejected_audit": "audit_site_visit",
            "pending_head": "sign_as_head",
            "pending_supervisor": "sign_as_supervisor",
            "pending_director": "sign_as_director",
        }
        can_act_now = False
        if item.status in stage_perm_map:
            can_act_now = await _has_perm(db, user, [stage_perm_map[item.status]])
        # Auditors can also view pending_head requests so they can re-audit
        # before the head signs.
        can_re_audit = (
            item.status == "pending_head"
            and item.head_signed_at is None
            and not item.head_signed_by
            and await _has_perm(db, user, ["audit_site_visit"])
        )
        if not (is_owner or can_view_all or can_act_now or can_re_audit):
            raise HTTPException(status_code=403, detail="ليس لديك صلاحية مشاهدة هذا الطلب")

    submitter = await _resolve_submitter_name(db, item.owner_id)
    return _serialize(item, submitted_by_name=submitter)


@router.post("/sign", response_model=SiteVisitRead)
async def sign_site_visit(
    data: SignSiteVisitRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    res = await db.execute(select(SiteVisitRequest).where(SiteVisitRequest.id == data.request_id))
    item = res.scalar_one_or_none()
    if not item:
        raise HTTPException(status_code=404, detail="الطلب غير موجود")

    # Build the textual signature from the current user's name (fallback to email,
    # then user id). This replaces the previous base64 image signature.
    signer_name = (user.get("name") or "").strip() or str(user.get("id") or "موقّع غير معروف")

    # Hard guard — requests still in the audit stage MUST NOT be signable.
    # The audit stage runs BEFORE the 3-stage signing chain, so any /sign
    # call against `pending_audit` or `rejected_audit` is rejected outright.
    if item.status in ("pending_audit", "rejected_audit"):
        raise HTTPException(
            status_code=400,
            detail="هذا الطلب بانتظار التدقيق ولا يمكن التوقيع عليه قبل اعتماده من المدقق",
        )

    now = _now()
    if item.status == "pending_head":
        if not await _has_perm(db, user, ["sign_as_head"]):
            raise HTTPException(status_code=403, detail="ليس لديك صلاحية التوقيع كرئيس قسم")
        item.head_signature = signer_name
        item.head_signed_at = now
        item.head_signed_by = user["id"]
        item.head_signed_by_name = signer_name
        item.status = "pending_supervisor"
        next_perm = "sign_as_supervisor"
        next_title = "طلب زيارة ميدانية بحاجة لتوقيع مراقب الصيانة"
        next_body = "تم توقيع رئيس القسم — بانتظار توقيعك"
    elif item.status == "pending_supervisor":
        if not await _has_perm(db, user, ["sign_as_supervisor"]):
            raise HTTPException(status_code=403, detail="ليس لديك صلاحية التوقيع كمراقب صيانة")
        item.supervisor_signature = signer_name
        item.supervisor_signed_at = now
        item.supervisor_signed_by = user["id"]
        item.supervisor_signed_by_name = signer_name
        item.status = "pending_director"
        next_perm = "sign_as_director"
        next_title = "طلب زيارة ميدانية بحاجة لتوقيع مدير الإدارة"
        next_body = "تم توقيع رئيس القسم ومراقب الصيانة — بانتظار توقيعك"
    elif item.status == "pending_director":
        if not await _has_perm(db, user, ["sign_as_director"]):
            raise HTTPException(status_code=403, detail="ليس لديك صلاحية التوقيع كمدير إدارة")
        item.director_signature = signer_name
        item.director_signed_at = now
        item.director_signed_by = user["id"]
        item.director_signed_by_name = signer_name
        item.status = "approved"
        next_perm = None
        next_title = "تم اعتماد طلب الزيارة الميدانية"
        next_body = "تمت الموافقة على طلبك بالكامل"
    else:
        raise HTTPException(status_code=400, detail="هذا الطلب لا يحتاج لتوقيع في هذه المرحلة")

    item.updated_at = now
    await db.commit()
    await db.refresh(item)
    logger.info(f"User {user['id']} signed site_visit_request {item.id} → {item.status}")

    # Notifications
    if next_perm:
        await _notify_users_with_perm(db, next_perm, next_title, next_body, item.id, exclude_user_id=user["id"])
    elif item.owner_id:
        # Final approval — notify owner
        try:
            from models.notifications import Notifications

            db.add(
                Notifications(
                    user_id=item.owner_id,
                    type="site_visit_request",
                    message=f"{next_title} — {next_body}",
                    report_id=0,
                    is_read=False,
                )
            )
            await db.commit()
        except Exception:
            await db.rollback()
        try:
            from services.web_push_service import send_push_to_users

            await send_push_to_users(
                db,
                [item.owner_id],
                next_title,
                next_body,
                notification_type="site_visit_request",
                url=f"/site-visit-requests/{item.id}",
            )
        except Exception:
            pass
        try:
            from services.ws_notifications import ws_notify_users

            await ws_notify_users(
                [item.owner_id],
                "site_visit_request",
                f"{next_title} — {next_body}",
                0,
                extra={"site_visit_id": item.id, "url": f"/site-visit-requests/{item.id}"},
            )
        except Exception:
            pass

    submitter = await _resolve_submitter_name(db, item.owner_id)
    return _serialize(item, submitted_by_name=submitter)


# ---------- Audit (gatekeeper before the signing chain) ----------
@router.post("/audit", response_model=SiteVisitRead)
async def audit_site_visit(
    data: AuditSiteVisitRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Approve or reject a site-visit request at the audit stage.

    Permission: caller must hold `audit_site_visit`.

    Allowed source statuses:
      - `pending_audit`   → may approve OR reject
      - `rejected_audit`  → may approve (re-open after submitter fixed it).
        Rejecting again is a no-op error to avoid status churn.
      - `pending_head`    → re-audit before head-of-department signs. The
        auditor may re-approve (refresh audit metadata + note) OR reject
        (send back to submitter as `rejected_audit`). Once the head has
        signed (head_signed_at IS NOT NULL or head_signed_by IS NOT NULL),
        re-audit is locked.

    On approve  → status becomes `pending_head` and head-of-department users
    are notified. On reject → status becomes `rejected_audit`, the auditor's
    note is stored in `audit_note`, and the submitter is notified so they
    can edit and resubmit.
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    if not await _has_perm(db, user, ["audit_site_visit"]):
        raise HTTPException(status_code=403, detail="ليس لديك صلاحية تدقيق طلبات الزيارات الميدانية")

    decision = (data.decision or "").strip().lower()
    if decision not in ("approve", "reject"):
        raise HTTPException(status_code=400, detail="قيمة القرار غير صحيحة")

    res = await db.execute(select(SiteVisitRequest).where(SiteVisitRequest.id == data.request_id))
    item = res.scalar_one_or_none()
    if not item:
        raise HTTPException(status_code=404, detail="الطلب غير موجود")

    # `pending_head` is allowed ONLY for re-audit, and ONLY when the head has
    # NOT yet signed. Once head_signed_at or head_signed_by is set, the
    # request has moved beyond the auditor's reach and re-audit is locked.
    if item.status == "pending_head":
        if item.head_signed_at is not None or item.head_signed_by:
            raise HTTPException(
                status_code=409,
                detail="تعذر إعادة التدقيق لأن رئيس القسم قد وقّع على الطلب",
            )
    elif item.status not in ("pending_audit", "rejected_audit"):
        raise HTTPException(
            status_code=400,
            detail="هذا الطلب ليس في مرحلة التدقيق",
        )
    if decision == "reject" and item.status == "rejected_audit":
        raise HTTPException(status_code=400, detail="الطلب مرفوض بالفعل")

    auditor_name = (user.get("name") or "").strip() or str(user.get("id") or "مدقق")
    now = _now()

    if decision == "approve":
        item.status = "pending_head"
        item.audited_by_id = user["id"]
        item.audited_by_name = auditor_name
        item.audited_at = now
        # Approval clears any previous rejection note so it doesn't leak
        # into the next stage.
        item.audit_note = None

        item.updated_at = now
        await db.commit()
        await db.refresh(item)
        logger.info(f"User {user['id']} APPROVED audit on site_visit_request {item.id}")

        # Notify head-of-department users
        try:
            await _notify_users_with_perm(
                db,
                "sign_as_head",
                "طلب زيارة ميدانية بحاجة لتوقيعك",
                f"تم اعتماد التدقيق — مقدم الطلب: {item.owner_name or ''}".strip(),
                item.id,
                exclude_user_id=user["id"],
            )
        except Exception:
            pass
    else:
        # Reject — note is mandatory so the submitter knows what to fix.
        note = (data.note or "").strip()
        if not note:
            raise HTTPException(status_code=400, detail="يجب إدخال سبب الرفض")

        item.status = "rejected_audit"
        item.audited_by_id = user["id"]
        item.audited_by_name = auditor_name
        item.audited_at = now
        item.audit_note = note

        item.updated_at = now
        await db.commit()
        await db.refresh(item)
        logger.info(f"User {user['id']} REJECTED audit on site_visit_request {item.id}")

        # Notify the submitter directly so they can fix and resubmit.
        if item.owner_id:
            title = "تم رفض طلب الزيارة الميدانية في التدقيق"
            body = f"السبب: {note}"
            try:
                from models.notifications import Notifications

                db.add(
                    Notifications(
                        user_id=item.owner_id,
                        type="site_visit_request",
                        message=f"{title} — {body}",
                        report_id=0,
                        is_read=False,
                    )
                )
                await db.commit()
            except Exception:
                await db.rollback()
            try:
                from services.web_push_service import send_push_to_users

                await send_push_to_users(
                    db,
                    [item.owner_id],
                    title,
                    body,
                    notification_type="site_visit_request",
                    url=f"/site-visit-requests/{item.id}",
                )
            except Exception:
                pass
            try:
                from services.ws_notifications import ws_notify_users

                await ws_notify_users(
                    [item.owner_id],
                    "site_visit_request",
                    f"{title} — {body}",
                    0,
                    extra={"site_visit_id": item.id, "url": f"/site-visit-requests/{item.id}"},
                )
            except Exception:
                pass

    submitter = await _resolve_submitter_name(db, item.owner_id)
    return _serialize(item, submitted_by_name=submitter)


# ---------- Update (re-submit after audit rejection) ----------
@router.post("/update", response_model=SiteVisitRead)
async def update_site_visit(
    data: CreateSiteVisitRequest,
    request_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Allow the original submitter to edit a request that was rejected by
    the auditor and re-submit it for audit, instead of creating a new one.

    Guards:
      - Caller must be the original `owner_id` of the request.
      - Current status MUST be ``rejected_audit``. Any other status (e.g.
        ``pending_audit``, ``pending_head``, ``approved``) is locked.

    On success:
      - All form fields are overwritten from the payload.
      - Stale audit metadata (``audited_by_id`` / ``audited_by_name`` /
        ``audited_at`` / ``audit_note``) is cleared.
      - ``status`` is set back to ``pending_audit`` so the request re-enters
        the audit queue.
      - ``edited_after_audit_at`` is set to ``now()`` so the auditor UI can
        show a "🔄 تم التعديل بعد الرفض" badge.
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    res = await db.execute(select(SiteVisitRequest).where(SiteVisitRequest.id == int(request_id)))
    item = res.scalar_one_or_none()
    if not item:
        raise HTTPException(status_code=404, detail="الطلب غير موجود")

    if item.owner_id != user.get("id"):
        raise HTTPException(status_code=403, detail="ليس لديك صلاحية تعديل هذا الطلب")
    if item.status != "rejected_audit":
        raise HTTPException(
            status_code=403,
            detail="لا يمكن تعديل الطلب في الحالة الحالية. التعديل متاح فقط بعد الرفض من التدقيق",
        )

    # Validate at least one row has data (date or description) — same rule as create.
    has_data = False
    for r in data.rows or []:
        if not isinstance(r, dict):
            continue
        if str(r.get("date") or "").strip() or str(r.get("description") or "").strip():
            has_data = True
            break
    if not has_data:
        raise HTTPException(status_code=400, detail="يجب تعبئة صف واحد على الأقل (التاريخ أو الوصف)")

    form_owner_name = (data.owner_name or "").strip() or None
    item.owner_name = form_owner_name or user.get("name")
    item.civil_id = (data.civil_id or "").strip() or None
    item.job_title = (data.job_title or "").strip() or None
    item.month = data.month
    item.year = data.year
    item.area = (data.area or "").strip() or None
    item.reason = (data.reason or "").strip() or None
    item.rows = json.dumps(data.rows or [], ensure_ascii=False)

    # If the form sent a fresh attendance_attachment value, accept it; otherwise
    # keep whatever was already on the record (the user may not have re-uploaded).
    new_att = (getattr(data, "attendance_attachment", None) or "").strip() or None
    if new_att:
        item.attendance_attachment = new_att

    # Clear stale audit metadata and re-queue for audit.
    item.audited_by_id = None
    item.audited_by_name = None
    item.audited_at = None
    item.audit_note = None
    item.status = "pending_audit"
    item.edited_after_audit_at = _now()
    item.updated_at = _now()

    await db.commit()
    await db.refresh(item)
    logger.info(f"User {user['id']} re-submitted site_visit_request {item.id} after audit rejection")

    # Notify auditors that a re-submission is awaiting their review.
    title = "طلب زيارة ميدانية مُعدَّل بحاجة لإعادة تدقيق"
    body = f"تم تعديل الطلب وإعادة إرساله — مقدم الطلب: {user.get('name') or user['id']}"
    await _notify_users_with_perm(db, "audit_site_visit", title, body, item.id, exclude_user_id=user["id"])

    submitter = (user.get("name") or "").strip() or None
    return _serialize(item, submitted_by_name=submitter)


# ---------- Attendance image upload ----------
_ALLOWED_ATTENDANCE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".gif"}
_MAX_ATTENDANCE_BYTES = 10 * 1024 * 1024  # 10 MB hard cap


def _attendance_dir() -> Path:
    """Return the on-disk directory for attendance attachments.

    The base uploads directory is whatever main.py resolved at startup and
    exported via the ``UPLOADS_DIR`` environment variable:
      - On Lambda (``IS_LAMBDA=true``) → ``/tmp/uploads`` (the only writable
        location on Lambda's read-only filesystem). Files there are
        ephemeral; they survive warm invocations on the same container but
        are wiped on cold starts and are NOT shared across concurrent
        containers. This unblocks the "تعذر حفظ المرفق" error in production
        without requiring an S3 / Atoms Cloud migration.
      - Off Lambda → ``<backend>/uploads`` as before.

    The directory (and the ``site-visit-attendance`` subdir) is mounted as
    ``/uploads/*`` by main.py's StaticFiles mount so files written here are
    served back via the same web path the DB stores.
    """
    base = os.environ.get("UPLOADS_DIR", "").strip()
    if base:
        d = Path(base) / "site-visit-attendance"
    else:
        # Fallback: same legacy path as before (used in tests / when main.py
        # didn't run, e.g. unit tests that import this router directly).
        backend_dir = Path(__file__).resolve().parent.parent
        d = backend_dir / "uploads" / "site-visit-attendance"
    d.mkdir(parents=True, exist_ok=True)
    return d


@router.post("/upload-attendance")
async def upload_attendance_image(
    request: Request,
    request_id: int,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """Upload (or replace) the attendance attachment image for a site visit
    request.

    Permissions:
      - The owner of the request can upload while it's still in
        ``pending_head`` (i.e. before the head signs).
      - Any user with ``view_all_site_visits`` can upload at ANY status
        (covers admins fixing missing attachments after the fact).

    Storage:
      - File is saved under ``<backend>/uploads/site-visit-attendance/`` with
        a random filename of the form ``req-{id}-{token}.{ext}`` so previous
        attachments are not overwritten by accident.
      - The previous attachment file (if any) is deleted on success.
      - The DB column ``attendance_attachment`` is set to the public web
        path (``/uploads/site-visit-attendance/<filename>``).
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    res = await db.execute(select(SiteVisitRequest).where(SiteVisitRequest.id == int(request_id)))
    item = res.scalar_one_or_none()
    if not item:
        raise HTTPException(status_code=404, detail="الطلب غير موجود")

    is_owner = item.owner_id == user.get("id")
    can_view_all = await _has_perm(db, user, ["view_all_site_visits"])
    # The owner may upload/replace the attendance image while the request is
    # still awaiting audit (`pending_audit`) or after the auditor rejected
    # it (`rejected_audit`). Once the request enters the signature chain
    # (`pending_head` and beyond) the owner can no longer change it. Admins
    # with `view_all_site_visits` can replace it at any stage.
    _OWNER_EDITABLE_STATUSES = ("pending_audit", "rejected_audit", "pending_head")
    if not can_view_all:
        if not is_owner:
            raise HTTPException(status_code=403, detail="ليس لديك صلاحية لرفع المرفق")
        if item.status not in _OWNER_EDITABLE_STATUSES:
            raise HTTPException(
                status_code=403,
                detail="لا يمكن تعديل المرفق بعد بدء عملية الاعتماد",
            )

    # Validate filename / extension
    original_name = (file.filename or "").strip()
    ext = Path(original_name).suffix.lower()
    if ext not in _ALLOWED_ATTENDANCE_EXTS:
        raise HTTPException(
            status_code=400,
            detail="صيغة الملف غير مدعومة. الصيغ المسموحة: jpg, jpeg, png, webp, gif",
        )

    # Read with size cap (stream-wise)
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="الملف فارغ")
    if len(content) > _MAX_ATTENDANCE_BYTES:
        raise HTTPException(status_code=413, detail="حجم الملف يتجاوز 10 ميجابايت")

    # Verify it's actually a valid image (not just a renamed file)
    try:
        from PIL import Image as _PILImage  # noqa: WPS433
        with _PILImage.open(io.BytesIO(content)) as _img:
            _img.verify()
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=400, detail="الملف ليس صورة صالحة") from exc

    # Save through the storage abstraction. On Lambda (where OSS is wired
    # up via OSS_SERVICE_URL + OSS_API_KEY) this uploads to Atoms Cloud OSS
    # and stores an ``oss://<bucket>/<key>`` URI in the DB. Off-Lambda or
    # without OSS credentials it falls back to the local
    # ``/uploads/site-visit-attendance/...`` disk layout. Either way, the
    # value returned here is what we persist on the request row.
    token = secrets.token_hex(8)
    new_name = f"req-{int(request_id)}-{token}{ext}"
    media_type = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp",
        ".gif": "image/gif",
    }.get(ext, "application/octet-stream")
    try:
        from services import attendance_storage as _att

        web_path = await _att.save_image(content, new_name, media_type)
    except Exception as exc:  # noqa: BLE001
        logger.error(f"Failed to save attendance image for req {request_id}: {exc}")
        raise HTTPException(status_code=500, detail="تعذر حفظ المرفق") from exc

    # Delete previous attachment (if any) — best-effort, don't fail if missing
    old_web_path = item.attendance_attachment
    if old_web_path:
        try:
            from services import attendance_storage as _att

            await _att.delete_image(old_web_path)
        except Exception as exc:  # noqa: BLE001
            logger.warning(f"Could not delete previous attendance image: {exc}")

    item.attendance_attachment = web_path
    item.updated_at = _now()
    # If the owner is replacing the attendance image AFTER the auditor
    # rejected the request, mark it as edited and re-queue for audit. This
    # mirrors the behaviour of POST /update so the auditor sees the
    # re-submission immediately even if only the image changed.
    if is_owner and item.status == "rejected_audit":
        item.edited_after_audit_at = _now()
        item.status = "pending_audit"
        # Clear stale audit metadata so the audit dialog shows a clean state.
        item.audited_by_id = None
        item.audited_by_name = None
        item.audited_at = None
        item.audit_note = None
    await db.commit()
    await db.refresh(item)
    logger.info(f"User {user['id']} uploaded attendance image for req {request_id}: {web_path}")

    return {
        "success": True,
        "request_id": int(request_id),
        "attendance_attachment": web_path,
    }


@router.post("/delete-attendance")
async def delete_attendance_image(
    payload: dict,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Remove the attendance attachment from a site visit request.

    Same permission rules as ``/upload-attendance``.
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    request_id = payload.get("request_id") if isinstance(payload, dict) else None
    if not request_id:
        raise HTTPException(status_code=400, detail="معرّف الطلب مطلوب")

    res = await db.execute(select(SiteVisitRequest).where(SiteVisitRequest.id == int(request_id))) 
    item = res.scalar_one_or_none()
    if not item:
        raise HTTPException(status_code=404, detail="الطلب غير موجود")

    is_owner = item.owner_id == user.get("id")
    can_view_all = await _has_perm(db, user, ["view_all_site_visits"])
    if not can_view_all:
        if not is_owner:
            raise HTTPException(status_code=403, detail="ليس لديك صلاحية لحذف المرفق")
        if item.status != "pending_head":
            raise HTTPException(
                status_code=403,
                detail="لا يمكن تعديل المرفق بعد بدء عملية الاعتماد",
            )

    web_path = item.attendance_attachment
    if web_path:
        try:
            from services import attendance_storage as _att

            await _att.delete_image(web_path)
        except Exception as exc:  # noqa: BLE001
            logger.warning(f"Could not delete attendance image file: {exc}")

    item.attendance_attachment = None
    item.updated_at = _now()
    await db.commit()
    logger.info(f"User {user['id']} deleted attendance image for req {request_id}")
    return {"success": True}


@router.get("/{request_id}/attendance-file")
async def get_attendance_file(
    request_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Serve the attendance image for a site visit request through an
    AUTHENTICATED API endpoint.

    Why this exists (production routing):
      - In dev, the Vite proxy forwards ``/uploads/*`` to the backend's
        ``StaticFiles`` mount.
      - On the deployed preview / production environment, the reverse proxy
        does NOT route ``/uploads/*`` to the backend (it's swallowed by the
        SPA fallback and returns HTTP 500). Browsing the static path directly
        therefore fails.
      - This API endpoint sidesteps that by serving the file through the
        regular ``/api/v1/...`` route, which IS routed to the backend.

    Permission model mirrors the rest of this router:
      - The owner of the request can always view their own attachment.
      - Anyone holding ``view_all_site_visits``, ``sign_as_head``,
        ``sign_as_supervisor``, ``sign_as_director``, or
        ``delete_site_visit`` may view it (these are the same roles that
        already see the request in the admin list / approval flow).
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    res = await db.execute(
        select(SiteVisitRequest).where(SiteVisitRequest.id == int(request_id))
    )
    item = res.scalar_one_or_none()
    if not item:
        raise HTTPException(status_code=404, detail="الطلب غير موجود")

    is_owner = item.owner_id == user.get("id")
    can_view = is_owner or await _has_perm(
        db,
        user,
        [
            "view_all_site_visits",
            "sign_as_head",
            "sign_as_supervisor",
            "sign_as_director",
            "delete_site_visit",
        ],
    )
    if not can_view:
        raise HTTPException(status_code=403, detail="ليس لديك صلاحية لعرض المرفق")

    web_path = (item.attendance_attachment or "").strip()
    if not web_path:
        raise HTTPException(status_code=404, detail="لا يوجد مرفق لهذا الطلب")

    # Resolve via the storage abstraction so this works for both
    # local-disk legacy paths AND new Atoms Cloud OSS paths
    # (``oss://<bucket>/<key>``). The helper returns raw bytes + media
    # type so we can stream them through this authenticated route
    # without ever exposing the underlying storage scheme to the client.
    try:
        from fastapi.responses import Response

        from services import attendance_storage as _att

        result = await _att.open_image_bytes(web_path)
    except Exception as exc:  # noqa: BLE001
        logger.warning(
            f"Attendance fetch failed for req {request_id} ({web_path!r}): {exc}"
        )
        raise HTTPException(status_code=500, detail="تعذر تحميل المرفق") from exc

    if not result:
        logger.warning(
            f"Attendance file missing for req {request_id}: {web_path!r}"
        )
        raise HTTPException(status_code=404, detail="الملف غير موجود على الخادم")

    content, media_type, file_name = result
    headers = {
        "Content-Disposition": f'inline; filename="{file_name or "attendance"}"',
        "Cache-Control": "private, max-age=60",
    }
    return Response(content=content, media_type=media_type, headers=headers)


@router.post("/delete")
async def delete_site_visit(
    payload: dict,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Delete a site visit request.

    Permission rule (per user requirement):
    - ONLY users with the dedicated `delete_site_visit` permission can delete a
      site-visit request, regardless of status or ownership.
    - Neither `view_all_site_visits` nor being the request owner grants
      deletion rights anymore.
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    request_id = payload.get("request_id") if isinstance(payload, dict) else None
    if not request_id:
        raise HTTPException(status_code=400, detail="معرّف الطلب مطلوب")

    res = await db.execute(select(SiteVisitRequest).where(SiteVisitRequest.id == int(request_id)))
    item = res.scalar_one_or_none()
    if not item:
        raise HTTPException(status_code=404, detail="الطلب غير موجود")

    can_delete = await _has_perm(db, user, ["delete_site_visit"])
    if not can_delete:
        raise HTTPException(
            status_code=403,
            detail="ليس لديك صلاحية حذف طلبات الزيارات الميدانية",
        )

    await db.delete(item)
    await db.commit()
    logger.info(f"User {user['id']} deleted site_visit_request {request_id}")
    return {"success": True, "deleted_id": int(request_id)}


# ---------- Helpers for the Word "site visit allowance" report ----------

_ARABIC_MONTH_NAMES = {
    1: "يناير", 2: "فبراير", 3: "مارس", 4: "أبريل",
    5: "مايو", 6: "يونيو", 7: "يوليو", 8: "أغسطس",
    9: "سبتمبر", 10: "أكتوبر", 11: "نوفمبر", 12: "ديسمبر",
}


def _arabic_month_name(m: int) -> str:
    """Return the Arabic name for a 1-12 month number, or the number as
    a string fallback for out-of-range values."""
    try:
        mi = int(m)
    except (TypeError, ValueError):
        return str(m or "")
    return _ARABIC_MONTH_NAMES.get(mi, str(mi))


def _extract_day_number(row: Dict[str, Any]) -> Optional[int]:
    """Extract just the day-of-month integer from a single visit row.

    Tries (in order): row['day'], row['extra']['day'], the leading number
    of row['date'] (e.g. "10/9/2025" → 10, "10-9" → 10, "10" → 10).
    Returns None when no usable day can be parsed.
    """
    if not isinstance(row, dict):
        return None
    candidates: List[Any] = [row.get("day")]
    extra = row.get("extra")
    if isinstance(extra, dict):
        candidates.append(extra.get("day"))
    date_v = row.get("date")
    if date_v:
        s = str(date_v).strip()
        # Take the first numeric token (handles "10/9", "10-9", "10/9/2025")
        token = ""
        for ch in s:
            if ch.isdigit():
                token += ch
            else:
                if token:
                    break
        if token:
            candidates.append(token)
    for c in candidates:
        if c is None:
            continue
        try:
            n = int(str(c).strip())
            if 1 <= n <= 31:
                return n
        except (TypeError, ValueError):
            continue
    return None


def _format_visit_days(rows: List[Dict[str, Any]]) -> str:
    """Build a "10/9/8" style string of unique visit days, sorted descending
    (most recent day first) to match the layout in the user's reference image.
    Empty rows / unparseable days are skipped."""
    days: List[int] = []
    seen = set()
    for r in rows or []:
        d = _extract_day_number(r)
        if d is None or d in seen:
            continue
        seen.add(d)
        days.append(d)
    days.sort(reverse=True)
    return "/".join(str(d) for d in days)


def _format_mosques(rows: List[Dict[str, Any]]) -> str:
    """Return a single label for the "اسم الموقع" column.

    Rules (per user request, image-1 (86).png):
      - 0 unique mosques → "" (empty cell)
      - 1 unique mosque  → that mosque's name (preserve accuracy)
      - 2+ unique mosques → "عدة مساجد" (Arabic for "multiple mosques")
    """
    seen: List[str] = []
    seen_set = set()
    for r in rows or []:
        if not isinstance(r, dict):
            continue
        m = str(r.get("mosque") or "").strip()
        if not m or m in seen_set:
            continue
        seen_set.add(m)
        seen.append(m)
    if not seen:
        return ""
    if len(seen) == 1:
        return seen[0]
    return "عدة مساجد"


def _count_non_empty_rows(rows: List[Dict[str, Any]]) -> int:
    """Count visit rows that contain at least one of: date/day, description,
    mosque. Used to fill the 'عدد الزيارات' column."""
    n = 0
    for r in rows or []:
        if not isinstance(r, dict):
            continue
        date_v = str(r.get("date") or "").strip()
        day_v = str(r.get("day") or "").strip()
        desc_v = str(r.get("description") or "").strip()
        mosque_v = str(r.get("mosque") or "").strip()
        extra = r.get("extra")
        if isinstance(extra, dict):
            day_v = day_v or str(extra.get("day") or "").strip()
        if date_v or day_v or desc_v or mosque_v:
            n += 1
    return n


def _set_cell_rtl(cell) -> None:
    """Force right-to-left direction on a docx table cell."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        # Remove existing bidi to avoid duplicates
        for old in tcPr.findall(qn("w:bidi")):
            tcPr.remove(old)
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        tcPr.append(bidi)
        for p in cell.paragraphs:
            pPr = p.paragraph_format.element.get_or_add_pPr()
            for old in pPr.findall(qn("w:bidi")):
                pPr.remove(old)
            b = OxmlElement("w:bidi")
            pPr.append(b)
    except Exception:
        pass


@router.post("/export-docx")
async def export_site_visits_docx(
    data: ExportDocxRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Export site-visit allowance requests for a month/year as a Word
    document that EXACTLY matches the official paper form layout:

    Heading (centered, bold, two lines):
        كشف مستحقي بدل موقع هندسي
        للعاملين بإدارة مساجد محافظة مبارك الكبير عن شهر <month_ar> لسنة <year> م

    Then a single 7-column RTL table:
        م | اسم المستلم | الرقم المدني | المسمى الوظيفي | اسم الموقع
        | تاريخ الأيام | عدد الزيارات

    One row per request. The "اسم الموقع" cell lists every unique mosque
    visited (comma-separated). The "تاريخ الأيام" cell lists every visit
    day as e.g. "10/9/8" (no months/years). The table is padded with empty
    rows up to a minimum of 10 to preserve the printed-form look.
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    if not await _has_perm(db, user, ["view_all_site_visits"]):
        raise HTTPException(status_code=403, detail="ليس لديك صلاحية تصدير الطلبات")

    # Defensive int casting in case the payload arrives with string month/year.
    try:
        month_i = int(data.month)
        year_i = int(data.year)
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="شهر/سنة غير صالحين")

    # Build the query. If `request_ids` is supplied (and non-empty), the
    # frontend is asking for an explicit pick of rows ("المحددة فقط") —
    # ignore month/year/status filters in that case.
    rid_list = [int(x) for x in (data.request_ids or []) if x is not None]
    if rid_list:
        q = select(SiteVisitRequest).where(SiteVisitRequest.id.in_(rid_list))
        q = q.order_by(SiteVisitRequest.id.asc())
        res = await db.execute(q)
        items = res.scalars().all()
        logger.info(
            f"export-docx: request_ids mode → {len(rid_list)} requested, "
            f"{len(items)} returned"
        )
        # Defense-in-depth: every explicitly-requested row must be fully
        # approved. Frontend already blocks this case with a toast, but a
        # malicious / outdated client could still POST request_ids of
        # non-approved rows — reject with a clear Arabic 400.
        not_approved = [r for r in items if r.status != "approved"]
        if not_approved:
            ids_str = "، ".join([f"#{r.id}" for r in not_approved])
            raise HTTPException(
                status_code=400,
                detail=(
                    "لا يمكن تصدير طلبات غير معتمدة. "
                    f"الطلبات التالية غير معتمدة: {ids_str}"
                ),
            )
    else:
        q = select(SiteVisitRequest).where(
            SiteVisitRequest.month == month_i,
            SiteVisitRequest.year == year_i,
        )
        # Scope precedence: an explicit named `scope` (e.g.
        # "audit_approved") overrides the legacy `include_unapproved` flag.
        # "audit_approved" = passed the audit stage = status NOT IN
        # ('pending_audit', 'rejected_audit'); i.e. every request currently
        # in or past the signature chain.
        if (data.scope or "").strip() == "audit_approved":
            q = q.where(
                SiteVisitRequest.status.notin_(["pending_audit", "rejected_audit"])
            )
        elif not data.include_unapproved:
            q = q.where(SiteVisitRequest.status == "approved")
        q = q.order_by(SiteVisitRequest.id.asc())
        res = await db.execute(q)
        items = res.scalars().all()
        logger.info(
            f"export-docx: month={month_i} year={year_i} "
            f"scope={data.scope!r} include_unapproved={data.include_unapproved} "
            f"→ {len(items)} items"
        )

    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENTATION
        from docx.shared import Pt, RGBColor, Cm, Mm
    except ImportError:
        logger.exception("python-docx missing — cannot generate Word file")
        raise HTTPException(
            status_code=500,
            detail="مكتبة python-docx غير مثبتة على الخادم — يرجى التواصل مع المسؤول",
        )

    doc = Document()

    # Default font + RTL for body text
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Switch the (only) section to A4 LANDSCAPE — gives "اسم الموقع" room
    # to fit several mosque names without wrapping the row to 4 lines.
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    # A4 landscape: 297mm × 210mm
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    month_ar = _arabic_month_name(month_i)

    # NOTE: Per user request, the formal cover letter (addressed to
    # "مدير إدارة الإسناد") is NO LONGER prepended here. It now lives in a
    # SEPARATE endpoint `/export-cover-letter-docx` which the frontend
    # exposes as its own export button. This endpoint produces ONLY the
    # data table (the "كشف مستحقي بدل موقع هندسي" sheet), so the user can
    # download the cover letter and the detailed sheet independently and
    # combine them as needed when printing/distributing.

    # ---- Two-line centered bold heading ----
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = h1.add_run("كشف مستحقي بدل موقع هندسي")
    r1.font.bold = True
    r1.font.size = Pt(18)
    r1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = h2.add_run(
        f"للعاملين بإدارة مساجد محافظة مبارك الكبير عن شهر {month_ar} لسنة {year_i} م"
    )
    r2.font.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # If no items, show only the heading + a small notice and STOP — do not
    # emit an empty table (per user request: row count must equal request count).
    if not items:
        doc.add_paragraph()  # spacer
        notice = doc.add_paragraph()
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = notice.add_run(f"لا توجد طلبات في {month_i}/{year_i}")
        nr.font.size = Pt(12)
        nr.font.bold = True
        nr.font.color.rgb = RGBColor(0x99, 0x33, 0x33)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"site-visits-{year_i}-{month_i:02d}.docx"
        return StreamingResponse(
            buf,
            media_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    doc.add_paragraph()  # spacer

    # ---- 7-column data table ----
    headers = [
        "م",
        "اسم المستلم",
        "الرقم المدني",
        "المسمى الوظيفي",
        "اسم الموقع",
        "تاريخ الأيام",
        "عدد الزيارات",
    ]

    # Pre-compute one display row per request. No padding — the number of data
    # rows MUST equal the number of actual requests (per user request).
    # `it.owner_name` is the name typed into the FORM by the submitter (not the
    # logged-in user's account username), which is what the printed sheet must
    # show.
    display_rows: List[List[str]] = []
    for idx, it in enumerate(items, start=1):
        try:
            parsed = json.loads(it.rows or "[]")
            row_list = [r for r in parsed if isinstance(r, dict)] if isinstance(parsed, list) else []
        except Exception:
            row_list = []

        display_rows.append([
            str(idx),
            (it.owner_name or "").strip() or "—",
            (it.civil_id or "").strip() or "—",
            (it.job_title or "").strip() or "—",
            _format_mosques(row_list) or "—",
            _format_visit_days(row_list) or "—",
            str(_count_non_empty_rows(row_list)),
        ])

    tbl = doc.add_table(rows=1 + len(display_rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        tbl.style = "Table Grid"  # all-sides borders
    except Exception:
        pass

    # Make the table itself RTL
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tblPr = tbl._tbl.tblPr
        bidi = OxmlElement("w:bidiVisual")
        bidi.set(qn("w:val"), "1")
        tblPr.append(bidi)
    except Exception:
        pass

    # Header row
    hdr_cells = tbl.rows[0].cells
    for ci, txt in enumerate(headers):
        cell = hdr_cells[ci]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p.add_run(txt)
        rn.font.bold = True
        rn.font.size = Pt(11)
        rn.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        _set_cell_rtl(cell)

    # Data rows
    for ri, vals in enumerate(display_rows, start=1):
        row_cells = tbl.rows[ri].cells
        for ci, v in enumerate(vals):
            cell = row_cells[ci]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rn = p.add_run(str(v))
            rn.font.size = Pt(11)
            rn.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            _set_cell_rtl(cell)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_size = buf.getbuffer().nbytes
    logger.info(f"export-docx: generated DOCX ({docx_size} bytes) for {len(items)} items")
    filename = f"site_visits_{year_i}_{month_i:02d}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(docx_size),
        },
    )


# ---------- Cover letter (standalone) ----------
class ExportCoverLetterRequest(BaseModel):
    """Payload for `/export-cover-letter-docx`.

    The cover letter is the formal one-page document addressed to
    "مدير إدارة الإسناد". It is independent from the data sheet and only
    needs a month/year — there are no rows, no request IDs, no status
    filters. The frontend collects month/year in its own dialog and posts
    them here. We accept both as `int` (defensive int casting handles
    string payloads from older clients).
    """
    month: int
    year: int


@router.post("/export-cover-letter-docx")
async def export_cover_letter_docx(
    data: ExportCoverLetterRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Generate ONLY the formal cover letter as a standalone Word document.

    Per user request, the cover letter that used to be page 1 of the main
    `/export-docx` output now lives in its own endpoint with its own
    button on the frontend. This way the user can:
      1. Download the data sheet (`/export-docx`) without the letter.
      2. Download the letter (`/export-cover-letter-docx`) on its own.
      3. Print/distribute either or both as needed.

    The output is a 1-page A4 LANDSCAPE document containing exactly the
    same blocks the cover page used to have:
      - Recipient line: "السيد / مدير إدارة الإسناد" (right) + "المحترم" (left)
      - Greeting "السلام عليكم ورحمة الله وبركاته،،،"
      - Subject "الموضوع : بدل موقع هندسي لشهر <month_ar> <year>"
        (centered, bold, underlined)
      - Two formal body paragraphs (also embedding the same month/year)
      - Closing line "وتفضلوا بقبول فائق التقدير والاحترام،"
      - Two-line signature block "مدير إدارة" / "مساجد محافظة مبارك الكبير"
      - Attachments header + bullet "- كشف التفصيلي والجداول"

    Permission: same as `/export-docx` — requires `view_all_site_visits`.
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    if not await _has_perm(db, user, ["view_all_site_visits"]):
        raise HTTPException(status_code=403, detail="ليس لديك صلاحية تصدير الطلبات")

    # Defensive int casting in case the payload arrives with strings.
    try:
        month_i = int(data.month)
        year_i = int(data.year)
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="شهر/سنة غير صالحين")
    if not (1 <= month_i <= 12):
        raise HTTPException(status_code=400, detail="الشهر يجب أن يكون بين 1 و 12")
    if not (2000 <= year_i <= 2100):
        raise HTTPException(status_code=400, detail="السنة غير صالحة")

    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENTATION
        from docx.shared import Pt, RGBColor, Cm, Mm
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlElement
    except ImportError:
        logger.exception("python-docx missing — cannot generate Word file")
        raise HTTPException(
            status_code=500,
            detail="مكتبة python-docx غير مثبتة على الخادم — يرجى التواصل مع المسؤول",
        )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Same section layout as the data-sheet export (A4 landscape) so when
    # the user prints both files together the orientation is consistent.
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    month_ar = _arabic_month_name(month_i)

    # Traditional Arabic = the classic, rounded calligraphic Arabic font
    # that matches the reference image (image-1 (102).png). On systems that
    # don't have it (Linux/LibreOffice), Word's complex-script shaper falls
    # back to whatever Arabic font is installed via the `cs` slot.
    _ARABIC_FONT = "Traditional Arabic"

    def _set_run_arabic_font(run, font_name: str) -> None:
        """Force `font_name` on all script slots so Word uses it for both
        Latin and Arabic (complex-script) glyphs."""
        try:
            rPr = run._element.get_or_add_rPr()
            for existing in rPr.findall(_qn("w:rFonts")):
                rPr.remove(existing)
            rFonts = _OxmlElement("w:rFonts")
            rFonts.set(_qn("w:ascii"), font_name)
            rFonts.set(_qn("w:hAnsi"), font_name)
            rFonts.set(_qn("w:cs"), font_name)
            rFonts.set(_qn("w:eastAsia"), font_name)
            rPr.insert(0, rFonts)
        except Exception:
            pass

    def _intro_para(text: str, *, size: int = 13, bold: bool = False,
                    underline: bool = False, align=WD_ALIGN_PARAGRAPH.RIGHT,
                    space_before: int = 0, space_after: int = 6) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        try:
            pPr = p._p.get_or_add_pPr()
            bidi = _OxmlElement("w:bidi")
            bidi.set(_qn("w:val"), "1")
            pPr.append(bidi)
        except Exception:
            pass
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = _ARABIC_FONT
        run.font.bold = bold
        run.font.underline = underline
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        try:
            rPr = run._element.get_or_add_rPr()
            szCs = _OxmlElement("w:szCs")
            szCs.set(_qn("w:val"), str(size * 2))
            rPr.append(szCs)
            if bold:
                bCs = _OxmlElement("w:bCs")
                bCs.set(_qn("w:val"), "1")
                rPr.append(bCs)
            rtl = _OxmlElement("w:rtl")
            rtl.set(_qn("w:val"), "1")
            rPr.append(rtl)
        except Exception:
            pass
        _set_run_arabic_font(run, _ARABIC_FONT)

    # Recipient address line — 1×2 borderless table so right cell holds the
    # title and left cell holds "المحترم" on the same line.
    addr_tbl = doc.add_table(rows=1, cols=2)
    addr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        tblPr = addr_tbl._tbl.tblPr
        bidi = _OxmlElement("w:bidiVisual")
        bidi.set(_qn("w:val"), "1")
        tblPr.append(bidi)
        tblW = _OxmlElement("w:tblW")
        tblW.set(_qn("w:w"), "5000")
        tblW.set(_qn("w:type"), "pct")
        tblPr.append(tblW)
    except Exception:
        pass

    addr_cells = addr_tbl.rows[0].cells
    rc = addr_cells[0]
    rc.text = ""
    pr = rc.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rrun = pr.add_run("السيد / مدير إدارة الإسناد")
    rrun.font.size = Pt(14)
    rrun.font.bold = True
    rrun.font.name = _ARABIC_FONT
    _set_run_arabic_font(rrun, _ARABIC_FONT)
    _set_cell_rtl(rc)

    lc = addr_cells[1]
    lc.text = ""
    pl = lc.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lrun = pl.add_run("المحترم")
    lrun.font.size = Pt(14)
    lrun.font.bold = True
    lrun.font.name = _ARABIC_FONT
    _set_run_arabic_font(lrun, _ARABIC_FONT)
    _set_cell_rtl(lc)

    # Strip borders — the table is purely a layout helper.
    try:
        for cell in (rc, lc):
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = _OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = _OxmlElement(f"w:{edge}")
                b.set(_qn("w:val"), "nil")
                tcBorders.append(b)
            tcPr.append(tcBorders)
    except Exception:
        pass

    # Greeting
    _intro_para(
        "السلام عليكم ورحمة الله وبركاته،،،",
        size=14,
        space_before=12,
        space_after=18,
    )

    # Subject — centered, bold, underlined
    _intro_para(
        f"الموضوع : بدل موقع هندسي لشهر {month_ar} {year_i}",
        size=16,
        bold=True,
        underline=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6,
        space_after=18,
    )

    # Body paragraph 1
    _intro_para(
        "تهديكم إدارة مساجد محافظة مبارك الكبير أطيب تحياتها وأمنياتها لكم بدوام الصحة والتوفيق،،",
        size=14,
        space_after=12,
    )

    # Body paragraph 2 — embeds the same month/year
    _intro_para(
        f"بالإشارة إلى الموضوع أعلاه، مرفق لسيادتكم كشوف بدل الموقع الهندسي (الأصلي) "
        f"لإدارة مساجد محافظة مبارك الكبير لشهر {month_ar} {year_i}م. برجاء من سيادتكم "
        "الإيعاز لمن يلزم لديكم بإرسال الكشف إلى السيد الوكيل المساعد لقطاع المساجد للتوقيع "
        "عليه ، ثم إرساله إلى إدارة الموارد البشرية لإصدار قرار الصرف الخاص به .",
        size=14,
        space_after=12,
    )

    # Closing
    _intro_para(
        "وتفضلوا بقبول فائق التقدير والاحترام،",
        size=14,
        space_after=24,
    )

    # Signature block
    _intro_para("مدير إدارة", size=13, bold=True, space_after=2)
    _intro_para(
        "مساجد محافظة مبارك الكبير",
        size=13,
        bold=True,
        space_after=24,
    )

    # Attachments
    _intro_para("المرفقات", size=14, bold=True, underline=True, space_after=4)
    _intro_para("- كشف التفصيلي والجداول", size=13, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_size = buf.getbuffer().nbytes
    logger.info(
        f"export-cover-letter-docx: generated DOCX ({docx_size} bytes) "
        f"for {month_i}/{year_i}"
    )
    filename = f"cover_letter_{year_i}_{month_i:02d}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(docx_size),
        },
    )


# ---------- Bulk sign ----------
@router.post("/bulk-sign")
async def bulk_sign_site_visits(
    data: BulkSignRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Sign multiple site-visit requests at once.

    For each request_id in the input list, the backend determines the current
    pending stage and applies the signer's textual name (full_name or email)
    if and only if the current user holds the matching stage permission AND
    the request is still at that stage. Returns per-request results.
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    if not data.request_ids:
        raise HTTPException(status_code=400, detail="يجب اختيار طلب واحد على الأقل")

    # Pre-compute permission flags once.
    can_head = await _has_perm(db, user, ["sign_as_head"])
    can_sup = await _has_perm(db, user, ["sign_as_supervisor"])
    can_dir = await _has_perm(db, user, ["sign_as_director"])
    if not (can_head or can_sup or can_dir):
        raise HTTPException(status_code=403, detail="ليس لديك أي صلاحية توقيع")

    signer_name = (user.get("name") or "").strip() or str(user.get("id") or "موقّع غير معروف")

    success_ids: List[int] = []
    failed: List[dict] = []
    notify_buckets: dict = {
        "sign_as_supervisor": [],  # next-stage items after head sign
        "sign_as_director": [],
    }
    final_approved_owner_ids: List[str] = []
    final_approved_request_ids: List[int] = []

    for rid in data.request_ids:
        try:
            r = await db.execute(select(SiteVisitRequest).where(SiteVisitRequest.id == int(rid)))
            item = r.scalar_one_or_none()
            if not item:
                failed.append({"id": rid, "error": "الطلب غير موجود"})
                continue
            now = _now()
            if item.status == "pending_head":
                if not can_head:
                    failed.append({"id": rid, "error": "لا تملك صلاحية التوقيع كرئيس قسم"})
                    continue
                item.head_signature = signer_name
                item.head_signed_at = now
                item.head_signed_by = user["id"]
                item.head_signed_by_name = signer_name
                item.status = "pending_supervisor"
                notify_buckets["sign_as_supervisor"].append(item.id)
            elif item.status == "pending_supervisor":
                if not can_sup:
                    failed.append({"id": rid, "error": "لا تملك صلاحية التوقيع كمراقب صيانة"})
                    continue
                item.supervisor_signature = signer_name
                item.supervisor_signed_at = now
                item.supervisor_signed_by = user["id"]
                item.supervisor_signed_by_name = signer_name
                item.status = "pending_director"
                notify_buckets["sign_as_director"].append(item.id)
            elif item.status == "pending_director":
                if not can_dir:
                    failed.append({"id": rid, "error": "لا تملك صلاحية التوقيع كمدير إدارة"})
                    continue
                item.director_signature = signer_name
                item.director_signed_at = now
                item.director_signed_by = user["id"]
                item.director_signed_by_name = signer_name
                item.status = "approved"
                if item.owner_id:
                    final_approved_owner_ids.append(item.owner_id)
                    final_approved_request_ids.append(item.id)
            else:
                failed.append({"id": rid, "error": "الطلب لا يحتاج لتوقيع في هذه المرحلة"})
                continue
            item.updated_at = now
            success_ids.append(item.id)
        except Exception as e:
            logger.exception(f"bulk-sign failed for {rid}: {e}")
            failed.append({"id": rid, "error": str(e)[:120]})

    if success_ids:
        try:
            await db.commit()
        except Exception as e:
            await db.rollback()
            logger.exception(f"bulk-sign commit failed: {e}")
            raise HTTPException(status_code=500, detail="تعذر حفظ التوقيعات الجماعية")

    # Best-effort notifications (don't block on failures).
    for perm_key, ids in notify_buckets.items():
        if not ids:
            continue
        for sid in ids:
            try:
                await _notify_users_with_perm(
                    db,
                    perm_key,
                    "طلب زيارة ميدانية بحاجة لتوقيعك",
                    f"بانتظار توقيعك على الطلب رقم #{sid}",
                    sid,
                    exclude_user_id=user["id"],
                )
            except Exception:
                pass

    if final_approved_owner_ids:
        try:
            from models.notifications import Notifications

            for oid, sid in zip(final_approved_owner_ids, final_approved_request_ids):
                db.add(
                    Notifications(
                        user_id=oid,
                        type="site_visit_request",
                        message=f"تم اعتماد طلب الزيارة الميدانية رقم #{sid}",
                        report_id=0,
                        is_read=False,
                    )
                )
            await db.commit()
        except Exception:
            await db.rollback()

    return {
        "success_count": len(success_ids),
        "success_ids": success_ids,
        "failed": failed,
    }


# ---------- Export approved requests as a ZIP of PDFs ----------
def _build_request_html(it: SiteVisitRequest, rows: List[dict]) -> str:
    """Render a single site-visit request as standalone HTML for PDF export."""
    status_labels = {
        "pending_head": "بانتظار رئيس القسم",
        "pending_supervisor": "بانتظار مراقب الصيانة",
        "pending_director": "بانتظار مدير الإدارة",
        "approved": "معتمد",
        "rejected": "مرفوض",
    }

    def esc(v: Any) -> str:
        if v is None:
            return ""
        s = str(v)
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )

    def fmt_dt(v: Any) -> str:
        if not v:
            return ""
        try:
            if isinstance(v, datetime):
                return v.strftime("%Y-%m-%d %H:%M")
            return str(v)
        except Exception:
            return str(v)

    rows_html_parts: List[str] = []
    for idx, r in enumerate(rows, start=1):
        if not isinstance(r, dict):
            continue
        rows_html_parts.append(
            "<tr>"
            f"<td>{idx}</td>"
            f"<td>{esc(r.get('date'))}</td>"
            f"<td>{esc(r.get('mosque'))}</td>"
            f"<td>{esc(r.get('description'))}</td>"
            f"<td>{esc(r.get('distance'))}</td>"
            f"<td>{esc(r.get('duration'))}</td>"
            "</tr>"
        )
    rows_html = (
        "".join(rows_html_parts)
        if rows_html_parts
        else '<tr><td colspan="6" style="text-align:center;color:#888">لا توجد زيارات مسجّلة</td></tr>'
    )

    return f"""<!DOCTYPE html>
<html dir="rtl" lang="ar"><head><meta charset="utf-8">
<title>طلب زيارة ميدانية #{it.id}</title>
<style>
@page {{ size: A4; margin: 14mm 12mm; }}
body {{ font-family: 'Noto Naskh Arabic', 'DejaVu Sans', sans-serif; font-size: 11pt; color:#111; direction: rtl; }}
h1 {{ font-size: 16pt; margin: 0 0 6mm 0; text-align:center; border-bottom:2px solid #1e40af; padding-bottom:4mm; color:#1e40af; }}
.meta {{ display: grid; grid-template-columns: repeat(2, 1fr); gap: 3mm 8mm; margin-bottom: 6mm; }}
.meta div {{ font-size: 10pt; }}
.meta b {{ display:block; color:#475569; font-size: 9pt; margin-bottom: 1mm; }}
table {{ width: 100%; border-collapse: collapse; margin-bottom: 5mm; }}
th, td {{ border: 1px solid #cbd5e1; padding: 2mm 2mm; text-align: right; font-size: 9.5pt; }}
th {{ background: #e0e7ff; color:#1e3a8a; }}
.sig-block {{ display:grid; grid-template-columns: repeat(3, 1fr); gap: 6mm; margin-top: 8mm; }}
.sig-card {{ border: 1px solid #cbd5e1; border-radius: 4px; padding: 3mm; text-align:center; min-height: 26mm; }}
.sig-card .role {{ font-weight: bold; font-size: 10pt; color:#1e40af; margin-bottom: 2mm; }}
.sig-card .name {{ font-size: 12pt; margin: 4mm 0 2mm 0; padding-bottom: 1mm; border-bottom: 1px solid #111; min-height: 6mm; }}
.sig-card .ts {{ font-size: 8pt; color:#64748b; }}
.sig-card.empty {{ background:#f8fafc; color:#94a3b8; }}
.notice {{ margin-top: 6mm; padding: 3mm; border:1px dashed #f59e0b; background:#fffbeb; color:#92400e; font-size:9pt; text-align:center; border-radius:3px; }}
.status {{ display:inline-block; padding: 1mm 3mm; border-radius:3px; font-size:9pt; }}
.status.approved {{ background:#dcfce7; color:#166534; }}
</style></head><body>
<h1>طلب زيارة ميدانية — بدل الموقع</h1>
<div class="meta">
  <div><b>رقم الطلب</b>#{it.id}</div>
  <div><b>الحالة</b><span class="status approved">{esc(status_labels.get(it.status, it.status))}</span></div>
  <div><b>مقدم الطلب</b>{esc(it.owner_name)}</div>
  <div><b>الرقم المدني</b>{esc(it.civil_id)}</div>
  <div><b>المسمى الوظيفي</b>{esc(it.job_title)}</div>
  <div><b>الشهر/السنة</b>{esc(it.month)}/{esc(it.year)}</div>
  <div><b>المنطقة</b>{esc(it.area)}</div>
  <div><b>السبب</b>{esc(it.reason)}</div>
</div>

<table>
<thead><tr>
<th style="width:8%">#</th>
<th style="width:14%">التاريخ</th>
<th style="width:22%">المسجد</th>
<th style="width:30%">الوصف</th>
<th style="width:13%">المسافة</th>
<th style="width:13%">المدة</th>
</tr></thead>
<tbody>{rows_html}</tbody>
</table>

<div class="sig-block">
  <div class="sig-card{(' empty' if not it.head_signature else '')}">
    <div class="role">رئيس القسم</div>
    <div class="name">{esc(it.head_signed_by_name) or '—'}</div>
    <div class="ts">{esc(fmt_dt(it.head_signed_at))}</div>
  </div>
  <div class="sig-card{(' empty' if not it.supervisor_signature else '')}">
    <div class="role">مراقب الصيانة</div>
    <div class="name">{esc(it.supervisor_signed_by_name) or '—'}</div>
    <div class="ts">{esc(fmt_dt(it.supervisor_signed_at))}</div>
  </div>
  <div class="sig-card{(' empty' if not it.director_signature else '')}">
    <div class="role">مدير الإدارة</div>
    <div class="name">{esc(it.director_signed_by_name) or '—'}</div>
    <div class="ts">{esc(fmt_dt(it.director_signed_at))}</div>
  </div>
</div>

<div class="notice">⚠️ يجب طباعة هذا النموذج وختمه يدوياً لإكمال الاعتماد الرسمي.</div>
</body></html>"""


def _generate_pdf_with_reportlab(it: SiteVisitRequest, rows: List[dict]) -> bytes:
    """Render a single SiteVisitRequest into PDF bytes using reportlab.

    reportlab is pure-Python (no native dependencies like weasyprint/cairo),
    so it works on any deployment target. We do simple text-mode rendering:
    a header, a metadata block, a visits table, and a 3-column signature
    footer. Arabic text is rendered via reportlab's stringWidth + drawString
    using a Unicode TTF font when available; otherwise we fall back to the
    built-in Helvetica which still renders the data (the user can still
    print the original HTML form when they need pixel-perfect Arabic).
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )

    # Try to register an Arabic-capable TTF if one is bundled in the system.
    # We probe the most common Linux paths; if none exist, we fall back to
    # Helvetica (Latin-only) — the data still renders, just not in shaped
    # Arabic. This is acceptable for an interim deployment-safe export.
    font_name = "Helvetica"
    bold_font_name = "Helvetica-Bold"
    candidate_fonts = [
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf", "/usr/share/fonts/truetype/noto/NotoSansArabic-Bold.ttf"),
        ("/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf", "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Bold.ttf"),
        ("/usr/share/fonts/truetype/amiri/Amiri-Regular.ttf", "/usr/share/fonts/truetype/amiri/Amiri-Bold.ttf"),
    ]
    import os as _os
    for reg_path, bold_path in candidate_fonts:
        if _os.path.exists(reg_path):
            try:
                pdfmetrics.registerFont(TTFont("ReportArabic", reg_path))
                font_name = "ReportArabic"
                if _os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont("ReportArabicBold", bold_path))
                    bold_font_name = "ReportArabicBold"
                else:
                    bold_font_name = "ReportArabic"
                break
            except Exception:
                continue

    status_labels = {
        "pending_head": "بانتظار رئيس القسم",
        "pending_supervisor": "بانتظار مراقب الصيانة",
        "pending_director": "بانتظار مدير الإدارة",
        "approved": "معتمد",
        "rejected": "مرفوض",
    }

    def fmt_dt(v: Any) -> str:
        if not v:
            return "—"
        try:
            if isinstance(v, datetime):
                return v.strftime("%Y-%m-%d %H:%M")
            return str(v)
        except Exception:
            return str(v)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=f"Site Visit Request #{it.id}",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "title_ar",
        parent=styles["Title"],
        fontName=bold_font_name,
        fontSize=16,
        alignment=1,  # CENTER
        textColor=colors.HexColor("#1e40af"),
        spaceAfter=10,
    )
    h2_style = ParagraphStyle(
        "h2_ar",
        parent=styles["Heading2"],
        fontName=bold_font_name,
        fontSize=11,
        textColor=colors.HexColor("#1e3a8a"),
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "body_ar",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9.5,
        leading=12,
    )

    story: list = []
    story.append(Paragraph(f"طلب زيارة ميدانية - بدل الموقع #{it.id}", title_style))
    story.append(Spacer(1, 4 * mm))

    # Metadata table (2 columns of label/value pairs)
    meta_rows = [
        ["رقم الطلب", f"#{it.id}", "الحالة", status_labels.get(it.status, it.status or "")],
        ["مقدم الطلب", it.owner_name or "—", "الرقم المدني", it.civil_id or "—"],
        ["المسمى الوظيفي", it.job_title or "—", "الشهر/السنة", f"{it.month or ''}/{it.year or ''}"],
        ["المنطقة", it.area or "—", "السبب", it.reason or "—"],
    ]
    meta_tbl = Table(meta_rows, colWidths=[28 * mm, 55 * mm, 28 * mm, 55 * mm])
    meta_tbl.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f1f5f9")),
                ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#f1f5f9")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#475569")),
                ("TEXTCOLOR", (2, 0), (2, -1), colors.HexColor("#475569")),
                ("FONTNAME", (0, 0), (0, -1), bold_font_name),
                ("FONTNAME", (2, 0), (2, -1), bold_font_name),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(meta_tbl)
    story.append(Spacer(1, 6 * mm))

    # Visits table
    story.append(Paragraph("الزيارات الميدانية", h2_style))
    visit_header = ["#", "التاريخ", "المسجد", "الوصف", "المسافة", "المدة"]
    visit_data: list = [visit_header]
    if rows:
        for idx, r in enumerate(rows, start=1):
            visit_data.append(
                [
                    str(idx),
                    str(r.get("date") or "—"),
                    str(r.get("mosque") or "—"),
                    str(r.get("description") or "—"),
                    str(r.get("distance") or "—"),
                    str(r.get("duration") or "—"),
                ]
            )
    else:
        visit_data.append(["—", "لا توجد زيارات مسجلة", "", "", "", ""])

    visits_tbl = Table(
        visit_data,
        colWidths=[10 * mm, 22 * mm, 38 * mm, 50 * mm, 26 * mm, 26 * mm],
        repeatRows=1,
    )
    visits_tbl.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, 0), bold_font_name),
                ("FONTNAME", (0, 1), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e0e7ff")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1e3a8a")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(visits_tbl)
    story.append(Spacer(1, 8 * mm))

    # Signatures footer (3 columns)
    story.append(Paragraph("التوقيعات", h2_style))
    sig_header = ["رئيس القسم", "مراقب الصيانة", "مدير الإدارة"]
    sig_names = [
        it.head_signed_by_name or it.head_signature or "—",
        it.supervisor_signed_by_name or it.supervisor_signature or "—",
        it.director_signed_by_name or it.director_signature or "—",
    ]
    sig_dates = [
        fmt_dt(it.head_signed_at),
        fmt_dt(it.supervisor_signed_at),
        fmt_dt(it.director_signed_at),
    ]
    sig_tbl = Table(
        [sig_header, sig_names, sig_dates],
        colWidths=[58 * mm, 58 * mm, 58 * mm],
    )
    sig_tbl.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, 0), bold_font_name),
                ("FONTNAME", (0, 1), (-1, 1), bold_font_name),
                ("FONTNAME", (0, 2), (-1, 2), font_name),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("FONTSIZE", (0, 1), (-1, 1), 11),
                ("FONTSIZE", (0, 2), (-1, 2), 8),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
                ("TEXTCOLOR", (0, 2), (-1, 2), colors.HexColor("#64748b")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 1), (-1, 1), 8),
                ("BOTTOMPADDING", (0, 1), (-1, 1), 8),
            ]
        )
    )
    story.append(sig_tbl)
    story.append(Spacer(1, 6 * mm))

    notice = Paragraph(
        "⚠️ يجب طباعة هذا النموذج وختمه يدوياً لإكمال الاعتماد الرسمي.",
        ParagraphStyle(
            "notice",
            parent=body_style,
            fontName=bold_font_name,
            alignment=1,
            textColor=colors.HexColor("#92400e"),
            backColor=colors.HexColor("#fffbeb"),
            borderColor=colors.HexColor("#f59e0b"),
            borderWidth=0.5,
            borderPadding=4,
        ),
    )
    story.append(notice)

    doc.build(story)
    pdf_bytes = buf.getvalue()
    buf.close()
    return pdf_bytes



# =====================================================================
# Playwright/Chromium PDF helper removed (Task 89) — superseded by
# ReportLab-based renderer in services/pdf_generator.py (Task 87).
# =====================================================================

@router.post("/export-ready-print-zip")
async def export_ready_print_zip(
    data: BatchReadyPrintRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Generate a ZIP of ready-to-print PDFs for a list of request IDs.

    Replaces the previous frontend pattern of looping `window.open(...)`
    per request — which was unreliable because browser popup-blockers
    typically suppress all but the first or last popup when more than
    2 windows are opened within a short interval.

    Each PDF is rendered via the same Playwright/Chromium pipeline used
    by `/export-approved-pdfs-zip`, so visit rows / approver names /
    signature images all render correctly with proper Arabic shaping.
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    if not data.request_ids:
        raise HTTPException(status_code=400, detail="يجب اختيار طلب واحد على الأقل")

    # ReportLab-based renderer — pure-Python, no native deps required.
    # Works on AWS Lambda's default Python runtime (no Pango/Cairo needed).
    # Lazy-import so a missing reportlab surfaces a clear Arabic message
    # at the right place rather than failing at module import time.
    try:
        from services.pdf_generator import (
            build_request_payload,
            render_site_visit_pdf,
        )
    except Exception as exc:  # noqa: BLE001
        logger.exception("PDF generator unavailable — cannot render PDFs")
        raise HTTPException(
            status_code=500,
            detail=(
                "محرّك توليد PDF غير متوفر على الخادم — "
                f"يرجى التواصل مع المسؤول. التفاصيل: {str(exc)[:200]}"
            ),
        )

    # Load all requested items in one pass; enforce per-item visibility rules.
    # Users who can view/sign at any stage should also be able to print —
    # otherwise signers (head/supervisor/director) without view_all_site_visits
    # cannot ready-print the requests they need to act on, which made the
    # button silently fail with "404 — محظور: N" for those users.
    can_view_all = await _has_perm(db, user, ["view_all_site_visits"])
    can_head = await _has_perm(db, user, ["sign_as_head"])
    can_sup = await _has_perm(db, user, ["sign_as_supervisor"])
    can_dir = await _has_perm(db, user, ["sign_as_director"])
    logger.info(
        f"[ready-print-zip] user={user.get('id')} role={user.get('role')} "
        f"view_all={can_view_all} head={can_head} sup={can_sup} dir={can_dir} "
        f"requested={data.request_ids}"
    )
    items: List[SiteVisitRequest] = []
    not_found: List[int] = []
    forbidden: List[int] = []
    for rid in data.request_ids:
        try:
            rid_int = int(rid)
        except (TypeError, ValueError):
            continue
        r = await db.execute(
            select(SiteVisitRequest).where(SiteVisitRequest.id == rid_int)
        )
        it = r.scalar_one_or_none()
        if not it:
            not_found.append(rid_int)
            continue
        # Permission: owner OR view_all OR a signer of any stage may print.
        is_owner = it.owner_id == user.get("id")
        is_signer = can_head or can_sup or can_dir
        if not (can_view_all or is_owner or is_signer):
            forbidden.append(rid_int)
            continue
        items.append(it)

    if not items:
        detail = "لم يتم العثور على أي طلب صالح للطباعة"
        if forbidden:
            detail += f" (محظور: {len(forbidden)})"
        if not_found:
            detail += f" (غير موجود: {len(not_found)})"
        raise HTTPException(status_code=404, detail=detail)

    import zipfile

    zip_buf = io.BytesIO()
    success_count = 0
    fail_count = 0
    failure_details: List[str] = []
    last_error: Optional[str] = None
    with zipfile.ZipFile(zip_buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for it in items:
            try:
                rows: List[dict] = []
                try:
                    parsed = json.loads(it.rows or "[]")
                    if isinstance(parsed, list):
                        rows = [r for r in parsed if isinstance(r, dict)]
                except Exception:
                    rows = []
                override_dict = (
                    data.override_names.model_dump() if data.override_names else None
                )
                payload = build_request_payload(
                    it, rows, override_names=override_dict
                )
                pdf_bytes = render_site_visit_pdf(payload)
                safe_owner = "".join(
                    c for c in (it.owner_name or "request") if c.isalnum() or c in (" ", "-", "_")
                ).strip().replace(" ", "_") or "request"
                fname = f"طلب-{it.id}-{safe_owner}.pdf"
                zf.writestr(fname, pdf_bytes)
                success_count += 1
            except Exception as e:
                fail_count += 1
                last_error = str(e)
                logger.exception(f"ready-print PDF failed for request {it.id}: {e}")
                failure_details.append(
                    f"- الطلب #{it.id} ({it.owner_name or '—'}): {str(e)[:300]}"
                )

    if success_count == 0 and fail_count > 0:
        detail = (
            f"تعذر توليد PDF لأي من الطلبات ({fail_count} فشل). "
            f"الخطأ: {(last_error or '')[:200]}"
        )
        raise HTTPException(status_code=500, detail=detail)

    if fail_count > 0:
        with zipfile.ZipFile(zip_buf, mode="a", compression=zipfile.ZIP_DEFLATED) as zf:
            report = (
                f"تقرير الطباعة الجماعية\n"
                f"=========================================\n"
                f"إجمالي المحدد: {len(data.request_ids)}\n"
                f"نجح: {success_count}\n"
                f"فشل: {fail_count}\n\n"
                f"تفاصيل الفشل:\n" + "\n".join(failure_details) + "\n"
            )
            zf.writestr("_تقرير_الطباعة.txt", report.encode("utf-8"))

    logger.info(
        f"[ready-print-zip] DONE: requested={len(data.request_ids)} "
        f"loaded={len(items)} success={success_count} failures={fail_count}"
    )
    zip_buf.seek(0)
    zip_size = zip_buf.getbuffer().nbytes
    filename = f"ready-print-{success_count}-requests.zip"
    return StreamingResponse(
        zip_buf,
        media_type="application/zip",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(zip_size),
        },
    )


@router.post("/export-approved-pdfs-zip")
async def export_approved_pdfs_zip(
    data: ExportPdfsZipRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Export site-visit requests for a given month/year as a single ZIP
    of PDFs (one per request).

    The PDFs are produced by rendering the **real** form HTML in a
    headless Chromium and exporting via `page.pdf()`. This avoids the
    Arabic-glyph rendering problems we hit with reportlab and ensures
    the PDF visually matches the printable form 1:1.
    """
    user = await _get_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="مطلوب تسجيل الدخول")

    if not await _has_perm(db, user, ["view_all_site_visits"]):
        raise HTTPException(status_code=403, detail="ليس لديك صلاحية تصدير الطلبات")

    # ReportLab-based renderer — pure-Python, no native deps. Lazy-import
    # so the router still imports even when reportlab is missing on a dev
    # box; we surface a clear Arabic 500 in that case.
    try:
        from services.pdf_generator import (
            build_request_payload,
            render_site_visit_pdf,
        )
    except Exception as exc:  # noqa: BLE001
        logger.exception("WeasyPrint generator unavailable — cannot render PDFs")
        raise HTTPException(
            status_code=500,
            detail=(
                "محرّك توليد PDF (WeasyPrint) غير متوفر على الخادم — "
                f"يرجى التواصل مع المسؤول. التفاصيل: {str(exc)[:200]}"
            ),
        )

    # Defensive int casting (in case the payload arrives with strings).
    try:
        month_i = int(data.month)
        year_i = int(data.year)
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="شهر/سنة غير صالحين")

    # Build the query. If `request_ids` is supplied (and non-empty), this is
    # the "المحددة فقط" mode — pick exactly those rows regardless of
    # month/year/status.
    rid_list = [int(x) for x in (data.request_ids or []) if x is not None]
    if rid_list:
        q = select(SiteVisitRequest).where(SiteVisitRequest.id.in_(rid_list))
        q = q.order_by(SiteVisitRequest.id.asc())
        res = await db.execute(q)
        items = res.scalars().all()
        logger.info(
            f"export-approved-pdfs-zip: request_ids mode → "
            f"{len(rid_list)} requested, {len(items)} returned"
        )
        if not items:
            raise HTTPException(
                status_code=404,
                detail="لم يتم العثور على أي من الطلبات المحددة",
            )
        # Defense-in-depth: every explicitly-requested row must be fully
        # approved. Same rule as /export-docx — see that endpoint for
        # rationale.
        not_approved = [r for r in items if r.status != "approved"]
        if not_approved:
            ids_str = "، ".join([f"#{r.id}" for r in not_approved])
            raise HTTPException(
                status_code=400,
                detail=(
                    "لا يمكن تصدير طلبات غير معتمدة. "
                    f"الطلبات التالية غير معتمدة: {ids_str}"
                ),
            )
    else:
        q = select(SiteVisitRequest).where(
            SiteVisitRequest.month == month_i,
            SiteVisitRequest.year == year_i,
        )
        # Scope precedence: an explicit named `scope` (e.g.
        # "audit_approved") overrides the legacy `include_unapproved` flag.
        # "audit_approved" = passed the audit stage = status NOT IN
        # ('pending_audit', 'rejected_audit'); i.e. every request currently
        # in or past the signature chain.
        if (data.scope or "").strip() == "audit_approved":
            q = q.where(
                SiteVisitRequest.status.notin_(["pending_audit", "rejected_audit"])
            )
        elif not data.include_unapproved:
            q = q.where(SiteVisitRequest.status == "approved")
        q = q.order_by(SiteVisitRequest.id.asc())
        res = await db.execute(q)
        items = res.scalars().all()

        logger.info(
            f"export-approved-pdfs-zip: month={month_i} year={year_i} "
            f"scope={data.scope!r} include_unapproved={data.include_unapproved} "
            f"→ found {len(items)} requests"
        )
        if not items:
            # Diagnostic: also count any requests in that month/year regardless of status,
            # so the user gets a clearer error message.
            total_q = select(SiteVisitRequest).where(
                SiteVisitRequest.month == month_i,
                SiteVisitRequest.year == year_i,
            )
            total_res = await db.execute(total_q)
            total_count = len(total_res.scalars().all())
            if total_count == 0:
                detail = f"لا توجد طلبات في {month_i}/{year_i}"
            else:
                detail = (
                    f"لا توجد طلبات معتمدة في {month_i}/{year_i} "
                    f"({total_count} طلب غير معتمد). فعّل خيار 'تضمين الطلبات غير المعتمدة' لتصدير الكل."
                )
            raise HTTPException(status_code=404, detail=detail)

    import zipfile

    zip_buf = io.BytesIO()
    success_count = 0
    fail_count = 0
    failure_details: List[str] = []
    last_error: Optional[str] = None
    with zipfile.ZipFile(zip_buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for it in items:
            try:
                rows: List[dict] = []
                try:
                    parsed = json.loads(it.rows or "[]")
                    if isinstance(parsed, list):
                        rows = [r for r in parsed if isinstance(r, dict)]
                except Exception:
                    rows = []
                payload = build_request_payload(it, rows)
                pdf_bytes = render_site_visit_pdf(payload)
                safe_owner = "".join(
                    c for c in (it.owner_name or "request") if c.isalnum() or c in (" ", "-", "_")
                ).strip().replace(" ", "_") or "request"
                # Status segment in filename so admins can tell at a glance which
                # requests are fully approved and which are still pending/rejected.
                status_label = {
                    "approved": "معتمد",
                    "pending_head": "بانتظار-رئيس-القسم",
                    "pending_supervisor": "بانتظار-المراقب",
                    "pending_director": "بانتظار-المدير",
                    "rejected": "مرفوض",
                }.get(it.status or "", "غير-معتمد")
                fname = f"site-visit-{it.id}-{safe_owner}-{status_label}.pdf"
                zf.writestr(fname, pdf_bytes)
                success_count += 1
            except Exception as e:
                fail_count += 1
                last_error = str(e)
                logger.exception(f"PDF generation failed for request {it.id}: {e}")
                failure_details.append(
                    f"- الطلب #{it.id} ({it.owner_name or '—'}): {str(e)[:300]}"
                )

    # If EVERYTHING failed, surface a clear 500 instead of returning a useless ZIP.
    if success_count == 0 and fail_count > 0:
        detail = (
            f"تعذر توليد PDF لأي من الطلبات ({fail_count} فشل). "
            f"الخطأ: {(last_error or '')[:200]}"
        )
        raise HTTPException(status_code=500, detail=detail)

    # If some failed, include a single consolidated report instead of one
    # FAILED-{id}.txt per failure (cleaner for the end user).
    if fail_count > 0:
        with zipfile.ZipFile(zip_buf, mode="a", compression=zipfile.ZIP_DEFLATED) as zf:
            report = (
                f"تقرير تصدير PDF — {month_i}/{year_i}\n"
                f"=========================================\n"
                f"إجمالي الطلبات: {len(items)}\n"
                f"نجح: {success_count}\n"
                f"فشل: {fail_count}\n\n"
                f"تفاصيل الفشل:\n" + "\n".join(failure_details) + "\n"
            )
            zf.writestr("_تقرير_التصدير.txt", report.encode("utf-8"))

    logger.info(
        f"export-approved-pdfs-zip: generated {success_count} PDFs, {fail_count} failures"
    )
    zip_buf.seek(0)
    zip_size = zip_buf.getbuffer().nbytes
    filename = f"site-visits-approved-{year_i}-{month_i:02d}.zip"
    return StreamingResponse(
        zip_buf,
        media_type="application/zip",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Length": str(zip_size),
        },
    )